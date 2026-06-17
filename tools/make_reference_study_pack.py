from __future__ import annotations

import csv
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "UWB_TurtleBot_교수님전달용_참고공부자료_그림추가"
DOCX_PATH = OUT_DIR / "01_참고자료_및_공부내용_교수님전달용.docx"
CSV_PATH = OUT_DIR / "02_참고자료_공부내용_매트릭스.csv"
MD_PATH = OUT_DIR / "03_공부내용_요약.md"
MSG_PATH = OUT_DIR / "04_교수님_전달메시지.txt"
README_PATH = OUT_DIR / "README_먼저보세요.txt"
ZIP_PATH = ROOT / "UWB_TurtleBot_교수님전달용_참고공부자료_그림추가.zip"
FIG_DIR = OUT_DIR / "05_논문그림자료"
FIG_LIST_PATH = OUT_DIR / "05_논문그림자료_목록.csv"
SOURCE_FIG_DIR = ROOT / "UWB_TurtleBot_교수님전달용_자료패키지" / "UWB_TurtleBot_교수님전달용_자료패키지" / "02_그림자료"


REFERENCES = [
    {
        "area": "프로젝트 요구사항",
        "title": "TurtleBot 자율 무선충전 프로젝트 가이드 / 팀 요구사항",
        "source": "수업/팀 내부 자료 및 본 repo의 docs/project_context.md",
        "url": "",
        "studied": "배터리 상태 확인부터 UWB 위치추정, 충전소 선택, LiDAR 접근, QR 정렬, 기계 가이드, 무선충전 시작까지의 전체 시나리오를 정리했다.",
        "applied": "우리 파트의 범위를 UWB 위치추정, 충전소 선택, LiDAR 기반 충전소 근방 접근, QR 도킹 handoff로 정의했다.",
    },
    {
        "area": "TurtleBot 플랫폼",
        "title": "ROBOTIS TurtleBot3 e-Manual: Waffle Pi 사양",
        "source": "ROBOTIS 공식 e-Manual",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/features/",
        "studied": "TurtleBot3 Waffle Pi가 Raspberry Pi, LDS LiDAR, Pi Camera를 포함하는 ROS 기반 모바일 로봇 플랫폼임을 확인했다.",
        "applied": "LiDAR는 장애물 회피/접근 주행, 카메라는 QR 최종 정렬, Raspberry Pi는 ROS 노드 실행 주체로 정리했다.",
    },
    {
        "area": "SLAM/Navigation",
        "title": "ROBOTIS TurtleBot3 Navigation 및 SLAM 문서",
        "source": "ROBOTIS 공식 e-Manual",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/navigation/",
        "studied": "일반 TurtleBot 내비게이션은 SLAM으로 map을 만든 뒤 그 map을 기반으로 goal까지 주행하는 구조임을 확인했다.",
        "applied": "본 프로젝트에서는 저장 map 기반 전역 내비게이션 대신 UWB 좌표와 LiDAR local planner를 이용한 충전소 접근 단계로 구현 범위를 좁혔다.",
    },
    {
        "area": "UWB 하드웨어",
        "title": "Qorvo/Decawave DW1000 자료",
        "source": "Qorvo DW1000 제품/데이터시트 및 DW1000 User Manual",
        "url": "https://www.qorvo.com/products/p/DW1000",
        "studied": "DW1000은 IEEE 802.15.4 UWB 기반 transceiver이며 two-way ranging 또는 TDoA 방식 위치추정에 활용될 수 있음을 공부했다.",
        "applied": "Arduino UWB tag/anchor 구성, ranging confirm/fail 로그 해석, anchor bias 및 거리 jump 보정 필요성의 근거로 사용했다.",
    },
    {
        "area": "UWB 위치추정",
        "title": "UWB 기반 WPT 단거리 정밀 위치추정 논문",
        "source": "PeerJ Computer Science / PMC 공개 논문",
        "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8176549/",
        "studied": "무선전력전송에서는 송신 코일과 수신 코일의 정렬 오차가 충전 효율과 성공 여부에 직접 영향을 준다는 점을 공부했다.",
        "applied": "UWB는 충전소 근방 접근과 위치 보조에 사용하고, 최종 10 cm급 정렬은 QR/카메라 및 기계 가이드가 이어받는 구조로 정리했다.",
    },
    {
        "area": "QR/비전 정렬",
        "title": "OpenCV QRCodeDetector 문서",
        "source": "OpenCV 공식 문서",
        "url": "https://docs.opencv.org/3.4/de/dc3/classcv_1_1QRCodeDetector.html",
        "studied": "카메라 영상에서 QR code를 검출하고, QR 중심점과 영상 중심점의 차이를 이용해 정렬 오차를 계산할 수 있음을 확인했다.",
        "applied": "UWB-LiDAR 접근 후 `/near_charger=true` 시점에 QR 기반 최종 정렬 단계로 넘기는 후속 구현 계획에 반영했다.",
    },
    {
        "area": "ROS 구현 구조",
        "title": "ROS topic 기반 센서/제어 분리",
        "source": "본 repo의 README.md, launch 파일, scripts 폴더",
        "url": "",
        "studied": "ROS에서 센서 입력과 제어 출력을 topic 단위로 나누면 UWB, LiDAR, charger target, velocity command를 독립적으로 검증할 수 있음을 정리했다.",
        "applied": "`/uwb_pose`, `/target_charger`, `/scan`, `/cmd_vel`, `/lidar_state`, `/near_charger` 구조로 노드를 구성했다.",
    },
    {
        "area": "실험/검증 방법",
        "title": "UWB 로그 분석 및 예비 검증 자료",
        "source": "uwb_logs_combined.csv, serial_logs/com5_live.log, config/*.yaml",
        "url": "",
        "studied": "UWB 좌표 로그에서 평균 위치, 표준편차, 기준점 대비 오차를 계산하고 COM5 시리얼 로그에서 ranging 안정성을 확인하는 방법을 정리했다.",
        "applied": "실제 로그 기반 예비 결과를 만들고, 최종 논문용 실험에서는 true_x/true_y 동기화 로그가 필요하다는 보완점을 도출했다.",
    },
]


FIGURES: list[dict[str, str]] = []


def get_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\malgunbd.ttf" if bold else r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill: str, width: int, line_gap: int = 6) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = "#cbd5e1", radius: int = 12) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=2)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#334155", width: int = 4) -> None:
    draw.line((start, end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 16 * direction, ey - 9), (ex - 16 * direction, ey + 9)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 9, ey - 16 * direction), (ex + 9, ey - 16 * direction)]
    draw.polygon(points, fill=color)


def base_canvas(title: str, subtitle: str = ""):
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(38, True)
    sub_font = get_font(22)
    draw.text((70, 46), title, font=title_font, fill="#0f172a")
    if subtitle:
        draw.text((72, 96), subtitle, font=sub_font, fill="#475569")
    return img, draw


def make_reference_to_system_figure(path: Path) -> None:
    img, draw = base_canvas("Reference Study to System Design Map", "공부한 자료가 프로젝트 설계로 연결되는 흐름")
    font_h = get_font(24, True)
    font = get_font(19)
    left = [
        ("TurtleBot3 e-Manual", "LiDAR, camera, Raspberry Pi, ROS platform"),
        ("DW1000 UWB 자료", "tag-anchor ranging, TWR/TDoA, calibration"),
        ("UWB-WPT 논문", "wireless charging needs short-range alignment"),
        ("OpenCV QR 문서", "QR detection and visual center offset"),
    ]
    right = [
        ("Platform Scope", "TurtleBot3 Waffle Pi + UWB tag + LiDAR + camera"),
        ("UWB Localization", "estimate pose and select nearest charger"),
        ("Local Approach", "LiDAR obstacle avoidance until 1 m radius"),
        ("Docking Handoff", "QR/camera alignment after near_charger=true"),
    ]
    for i, (h, body) in enumerate(left):
        y = 170 + i * 160
        rounded_box(draw, (70, y, 610, y + 105), "#f8fafc")
        draw.text((95, y + 18), h, font=font_h, fill="#1e3a8a")
        draw.text((95, y + 56), body, font=font, fill="#334155")
    for i, (h, body) in enumerate(right):
        y = 170 + i * 160
        rounded_box(draw, (990, y, 1530, y + 105), "#eef6ff")
        draw.text((1015, y + 18), h, font=font_h, fill="#0f766e")
        draw.text((1015, y + 56), body, font=font, fill="#334155")
        arrow(draw, (620, y + 52), (980, y + 52), "#64748b")
    rounded_box(draw, (665, 300, 935, 560), "#fff7ed", "#fed7aa")
    draw.text((700, 330), "Our Design", font=get_font(30, True), fill="#9a3412")
    draw_wrapped(draw, (700, 385), "UWB-LiDAR approach first, QR docking later", get_font(24, True), "#431407", 200)
    img.save(path)


def make_uwb_geometry_figure(path: Path) -> None:
    img, draw = base_canvas("UWB Anchor Geometry and Position Estimation", "4-anchor ranging concept used for TurtleBot pose estimation")
    font = get_font(22)
    small = get_font(18)
    field = (220, 170, 1120, 820)
    draw.rectangle(field, outline="#94a3b8", width=4)
    anchors = {"A1": (220, 820), "A2": (1120, 820), "A3": (220, 170), "A4": (1120, 170)}
    robot = (650, 520)
    for name, pos in anchors.items():
        x, y = pos
        draw.ellipse((x - 22, y - 22, x + 22, y + 22), fill="#1d4ed8")
        draw.text((x - 26, y - 62 if y > 500 else y + 34), name, font=font, fill="#1e3a8a")
        draw.line((x, y, robot[0], robot[1]), fill="#bfdbfe", width=3)
        mid = ((x + robot[0]) // 2, (y + robot[1]) // 2)
        draw.text((mid[0] + 8, mid[1] - 12), "r_i", font=small, fill="#2563eb")
    draw.ellipse((robot[0] - 34, robot[1] - 34, robot[0] + 34, robot[1] + 34), fill="#dc2626")
    draw.text((robot[0] - 92, robot[1] + 46), "TurtleBot tag (x, y)", font=font, fill="#991b1b")
    rounded_box(draw, (1180, 240, 1510, 690), "#f8fafc")
    draw.text((1210, 270), "Studied", font=get_font(28, True), fill="#0f172a")
    bullets = [
        "distance from tag to each anchor",
        "anchor coordinate accuracy matters",
        "bias/jump rejection improves stability",
        "pose is used for charger selection",
    ]
    y = 325
    for bullet in bullets:
        draw.text((1210, y), f"- {bullet}", font=small, fill="#334155")
        y += 54
    img.save(path)


def make_ros_topic_figure(path: Path) -> None:
    img, draw = base_canvas("ROS Topic Pipeline for UWB-LiDAR Charging Approach", "센서 입력, 목표 선택, 주행 출력, QR handoff를 topic으로 분리")
    boxes = [
        ((70, 210, 390, 330), "UWB ranges / pose", "/uwb/ranges\n/uwb_pose"),
        ((470, 210, 790, 330), "Charger selector", "/selected_charger_id\n/target_charger"),
        ((870, 210, 1190, 330), "LiDAR local planner", "/scan + /uwb_pose\nheading + obstacle"),
        ((1270, 210, 1530, 330), "Robot command", "/cmd_vel\n/lidar_state"),
        ((870, 520, 1190, 640), "Near charger", "/near_charger=true\nstop radius 1.0 m"),
        ((1270, 520, 1530, 640), "Future docking", "QR detection\ncamera alignment"),
    ]
    for box, h, body in boxes:
        rounded_box(draw, box, "#f8fafc")
        draw.text((box[0] + 20, box[1] + 22), h, font=get_font(24, True), fill="#1e3a8a")
        draw_wrapped(draw, (box[0] + 20, box[1] + 60), body, get_font(20), "#334155", box[2] - box[0] - 40)
    arrow(draw, (390, 270), (470, 270))
    arrow(draw, (790, 270), (870, 270))
    arrow(draw, (1190, 270), (1270, 270))
    arrow(draw, (1030, 330), (1030, 520))
    arrow(draw, (1190, 580), (1270, 580))
    rounded_box(draw, (145, 710, 1455, 820), "#ecfeff", "#67e8f9")
    draw.text((180, 738), "Key point for the paper:", font=get_font(26, True), fill="#155e75")
    draw_wrapped(
        draw,
        (520, 733),
        "UWB provides charger direction, LiDAR handles local obstacles, and QR handles final alignment.",
        get_font(22),
        "#164e63",
        850,
    )
    img.save(path)


def make_paper_figure_index(path: Path) -> None:
    img, draw = base_canvas("Recommended Figures for Paper / Report", "본문에 넣을 그림과 아직 촬영해야 할 사진을 구분")
    rows = [
        ("Fig. 1", "System architecture", "already prepared"),
        ("Fig. 2", "Operation flow", "already prepared"),
        ("Fig. 3", "Experiment layout", "already prepared"),
        ("Fig. 4", "Reference-study map", "new diagram"),
        ("Fig. 5", "UWB anchor geometry", "new diagram"),
        ("Fig. 6", "ROS topic pipeline", "new diagram"),
        ("Fig. 7", "UWB static scatter/error graphs", "actual log based"),
        ("Photo A", "Real robot + UWB tag mounting", "take in lab"),
        ("Photo B", "Charger/QR marker setup", "take in lab"),
    ]
    x0, y0 = 90, 170
    col = [140, 780, 320]
    headers = ["Item", "Content", "Status"]
    x = x0
    for i, header in enumerate(headers):
        draw.rectangle((x, y0, x + col[i], y0 + 58), fill="#e8eef5", outline="#cbd5e1", width=2)
        draw.text((x + 18, y0 + 16), header, font=get_font(22, True), fill="#0f172a")
        x += col[i]
    y = y0 + 58
    for item, content, status in rows:
        vals = [item, content, status]
        x = x0
        for i, val in enumerate(vals):
            draw.rectangle((x, y, x + col[i], y + 62), fill="white", outline="#cbd5e1", width=2)
            draw.text((x + 18, y + 18), val, font=get_font(20), fill="#334155")
            x += col[i]
        y += 62
    img.save(path)


def copy_existing_figures() -> list[dict[str, str]]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    copied = []
    source_map = [
        ("fig01_system_architecture.png", "fig1_system_architecture.png", "전체 시스템 구성도", "논문 방법론: 시스템 구성"),
        ("fig02_operation_flow.png", "fig2_operation_flow.png", "자율 충전 동작 흐름도", "논문 방법론: 알고리즘 흐름"),
        ("fig03_experiment_layout.png", "fig3_experiment_layout.png", "UWB/충전소 실험 배치도", "논문 실험환경"),
        ("fig07_uwb_static_scatter.png", "fig4_uwb_static_scatter_20260513.png", "실제 UWB 정적 좌표 산점도", "논문 예비 결과"),
        ("fig08_uwb_static_error.png", "fig5_uwb_static_error_20260513.png", "실제 UWB 기준점 오차 그래프", "논문 예비 결과"),
        ("fig09_all_uwb_log_trace.png", "fig6_all_uwb_logs_trace_20260513.png", "저장된 UWB 로그 궤적", "논문 예비 결과/부록"),
    ]
    for dest_name, source_name, title, use in source_map:
        source = SOURCE_FIG_DIR / source_name
        dest = FIG_DIR / dest_name
        if source.exists():
            shutil.copy2(source, dest)
            copied.append({"file": dest.name, "title": title, "use": use, "note": "기존 패키지 그림/실제 로그 그래프"})
    return copied


def create_figure_assets() -> list[dict[str, str]]:
    figures = copy_existing_figures()
    generated = [
        ("fig04_reference_to_system_map.png", "참고자료-시스템 설계 연결도", "논문 서론/방법론", make_reference_to_system_figure),
        ("fig05_uwb_anchor_geometry.png", "UWB anchor geometry 및 위치추정 개념도", "논문 UWB 위치추정 설명", make_uwb_geometry_figure),
        ("fig06_ros_topic_pipeline.png", "ROS topic 기반 UWB-LiDAR 접근 pipeline", "논문 구현 구조", make_ros_topic_figure),
        ("fig10_paper_figure_index.png", "논문 삽입 그림/사진 후보 목록", "교수님 논문 구성 협의", make_paper_figure_index),
    ]
    for file_name, title, use, maker in generated:
        path = FIG_DIR / file_name
        maker(path)
        figures.append({"file": file_name, "title": title, "use": use, "note": "새로 생성한 논문용 개념도"})
    figures.sort(key=lambda item: item["file"])
    return figures


def write_figure_list(figures: list[dict[str, str]]) -> None:
    with FIG_LIST_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["file", "title", "paper_use", "note"])
        for fig in figures:
            writer.writerow([fig["file"], fig["title"], fig["use"], fig["note"]])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 15, "1F4D78", 12, 6),
        ("Heading 2", 12, "2E74B5", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row.cells[i])

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True

    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.text = value
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.6) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def build_docx(figures: list[dict[str, str]]) -> None:
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("UWB TurtleBot 자율 무선충전 프로젝트")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    run = p.add_run("참고자료 및 공부내용 정리 | 교수님 전달용 | 2026-05-31")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. 전달 목적", level=1)
    doc.add_paragraph(
        "본 문서는 실험 결과 보고서가 아니라, UWB 기반 TurtleBot 자율 무선충전 프로젝트를 진행하면서 "
        "팀이 어떤 자료를 참고했고 각 자료에서 무엇을 공부했으며, 그 내용을 프로젝트 설계와 구현에 어떻게 반영했는지 정리한 자료이다."
    )
    add_table(
        doc,
        ["구분", "정리 내용"],
        [
            ["핵심 질문", "우리가 무엇을 참고했고, 무엇을 공부했으며, 프로젝트에 어떻게 적용했는가"],
            ["현재 구현 범위", "UWB 위치추정, 충전소 선택, LiDAR local approach, QR 도킹 handoff"],
            ["아직 분리한 범위", "QR 최종 정렬 실측, 기계 가이드, 실제 무선충전 시작 성공률"],
        ],
        [1.45, 5.55],
    )

    doc.add_heading("2. 참고자료별 공부내용", level=1)
    add_table(
        doc,
        ["분야", "참고자료", "공부한 내용", "프로젝트 반영"],
        [
            [item["area"], item["title"], item["studied"], item["applied"]]
            for item in REFERENCES
        ],
        [0.9, 1.75, 2.15, 2.2],
    )

    doc.add_heading("3. 논문/보고서 삽입 후보 그림", level=1)
    doc.add_paragraph(
        "아래 그림들은 교수님께 논문 방향을 설명할 때 바로 사용할 수 있도록 정리한 후보 이미지이다. "
        "시스템 구성도, 동작 흐름도, UWB 좌표계, 실제 로그 기반 그래프, 참고자료-구현 연결도를 함께 포함했다. "
        "실제 로봇 장착 사진과 충전기/QR 마커 사진은 현장에서 촬영한 뒤 추가하면 된다."
    )
    add_table(
        doc,
        ["그림 파일", "내용", "논문 내 사용 위치"],
        [[fig["file"], fig["title"], fig["use"]] for fig in figures],
        [1.85, 2.8, 2.35],
    )
    for fig in figures:
        add_figure(doc, FIG_DIR / fig["file"], f"{fig['file']} - {fig['title']}", 6.4)

    doc.add_heading("4. 공부한 내용을 프로젝트 구조로 연결", level=1)
    add_table(
        doc,
        ["공부한 개념", "프로젝트에서의 연결"],
        [
            ["UWB tag-anchor ranging", "로봇의 현재 위치를 UWB 좌표계에서 추정하고 충전소 선택에 사용"],
            ["Anchor geometry와 calibration", "UWB 오차를 줄이기 위해 anchor 좌표, bias, jump rejection, smoothing을 설정"],
            ["TurtleBot3 LiDAR/Camera 구성", "LiDAR는 접근/장애물 회피, 카메라는 QR 최종 정렬 단계에 배치"],
            ["SLAM navigation 구조", "일반 map 기반 navigation과 달리 본 구현은 UWB pose 기반 mapless local approach로 제한"],
            ["WPT 정렬 오차 중요성", "UWB만으로 최종 충전 위치를 맞추기보다 QR/기계 가이드 handoff가 필요하다고 판단"],
            ["ROS topic 분리", "각 기능을 `/uwb_pose`, `/target_charger`, `/scan`, `/cmd_vel`, `/near_charger`로 나누어 검증"],
        ],
        [2.0, 5.0],
    )

    doc.add_heading("5. 교수님께 설명할 핵심 요약", level=1)
    doc.add_paragraph(
        "저희는 TurtleBot3 공식 자료를 통해 플랫폼의 LiDAR/카메라/ROS 구성을 확인했고, DW1000 UWB 자료를 통해 "
        "tag-anchor 거리 측정과 보정 필요성을 공부했습니다. 또한 UWB 기반 WPT 정밀 위치추정 논문을 참고해 "
        "무선충전에서는 단순 접근뿐 아니라 최종 정렬 오차가 중요하다는 점을 이해했습니다. 이를 바탕으로 현재 구현은 "
        "UWB 위치추정과 충전소 선택, LiDAR 기반 근접 접근, QR 기반 최종 도킹 handoff 구조로 정리했습니다."
    )

    doc.add_heading("6. 참고 링크", level=1)
    for item in REFERENCES:
        if item["url"]:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.add_run(f"{item['title']}: ").bold = True
            p.add_run(item["url"])

    doc.save(DOCX_PATH)


def write_csv() -> None:
    with CSV_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["분야", "참고자료", "출처", "URL", "공부한 내용", "프로젝트 반영"])
        for item in REFERENCES:
            writer.writerow([item["area"], item["title"], item["source"], item["url"], item["studied"], item["applied"]])


def write_markdown() -> None:
    lines = [
        "# UWB TurtleBot 참고자료 및 공부내용 요약",
        "",
        "## 한 문장 요약",
        "",
        "TurtleBot3 공식 문서, DW1000 UWB 자료, UWB 기반 WPT 정밀 위치추정 논문, OpenCV QR 문서를 참고하여 UWB 위치추정-충전소 선택-LiDAR 접근-QR 도킹 handoff 구조를 설계했다.",
        "",
        "## 참고자료별 정리",
        "",
    ]
    for item in REFERENCES:
        lines += [
            f"### {item['area']} - {item['title']}",
            f"- 출처: {item['source']}",
            f"- URL: {item['url'] or '내부/수업 자료'}",
            f"- 공부한 내용: {item['studied']}",
            f"- 프로젝트 반영: {item['applied']}",
            "",
        ]
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def write_message() -> None:
    MSG_PATH.write_text(
        """교수님, 안녕하세요.

요청하신 대로 저희가 UWB TurtleBot 자율 무선충전 프로젝트를 진행하면서 참고한 자료와 공부한 내용을 따로 정리해 전달드립니다.

자료는 단순 참고문헌 목록이 아니라, 각 자료에서 무엇을 공부했고 그 내용을 프로젝트의 어느 부분에 반영했는지 중심으로 구성했습니다. 주요 참고 범위는 TurtleBot3 공식 e-Manual, DW1000 UWB 자료, UWB 기반 WPT 정밀 위치추정 논문, OpenCV QR 검출 문서, 그리고 저희 ROS 구현 구조입니다.

추가로 논문/보고서에 바로 넣을 수 있는 그림 후보도 함께 넣었습니다. 시스템 구성도, 동작 흐름도, UWB anchor geometry 개념도, ROS topic pipeline, 실제 UWB 로그 기반 그래프를 포함했습니다. 실제 로봇 장착 사진과 충전기/QR 마커 사진은 현장 촬영 후 추가하면 됩니다.

현재 저희 구현 범위는 UWB 위치추정, 충전소 선택, LiDAR 기반 충전소 근방 접근, QR 도킹 단계로의 handoff이며, QR 최종 정렬과 실제 무선충전 시작 성공률은 후속 실험 범위로 분리했습니다.

감사합니다.
""",
        encoding="utf-8",
    )

    README_PATH.write_text(
        """이 폴더는 교수님께 전달할 '참고자료 및 공부내용' 전용 패키지입니다.

먼저 볼 파일:
1. 01_참고자료_및_공부내용_교수님전달용.docx
2. 04_교수님_전달메시지.txt

보조 파일:
- 02_참고자료_공부내용_매트릭스.csv: 참고자료별 공부내용/프로젝트 반영 표
- 03_공부내용_요약.md: 같은 내용을 markdown으로 정리한 파일
- 05_논문그림자료/: 논문/보고서 삽입 후보 그림 및 실제 로그 기반 그래프
- 05_논문그림자료_목록.csv: 그림별 사용 위치 정리
""",
        encoding="utf-8",
    )


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figures = create_figure_assets()
    write_figure_list(figures)
    build_docx(figures)
    write_csv()
    write_markdown()
    write_message()
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    shutil.make_archive(str(ZIP_PATH.with_suffix("")), "zip", OUT_DIR)
    print(f"Wrote {OUT_DIR}")
    print(f"Wrote {ZIP_PATH}")


if __name__ == "__main__":
    main()
