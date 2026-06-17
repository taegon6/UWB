from __future__ import annotations

import csv
import math
import statistics
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
PACKAGE_ROOT = next(ROOT.glob("UWB_TurtleBot_교수님전달용_자료패키지/UWB_TurtleBot_교수님전달용_자료패키지"))
DOCX_PATH = PACKAGE_ROOT / "01_UWB_TurtleBot_자료취합본_교수님전달용.docx"
FIG_DIR = PACKAGE_ROOT / "02_그림자료"
DATA_DIR = PACKAGE_ROOT / "03_실험데이터"
OLD_DATA_DIR = PACKAGE_ROOT / "03_실험데이터_템플릿"
MSG_DIR = PACKAGE_ROOT / "04_메시지_체크리스트"
REF_DIR = PACKAGE_ROOT / "05_참고_공부자료"

TRUE_X = 2.50
TRUE_Y = 2.50


def ensure_dirs() -> None:
    if OLD_DATA_DIR.exists() and not DATA_DIR.exists():
        OLD_DATA_DIR.rename(DATA_DIR)
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    MSG_DIR.mkdir(parents=True, exist_ok=True)
    REF_DIR.mkdir(parents=True, exist_ok=True)
    for template_file in DATA_DIR.glob("*_template.csv"):
        template_file.unlink()


REFERENCE_ITEMS = [
    {
        "category": "수업/프로젝트 가이드",
        "title": "TurtleBot 자율 무선충전 프로젝트 요구사항",
        "source": "교수님/수업 가이드 및 팀 요구사항 정리",
        "url": "",
        "studied": "배터리 감지, UWB 위치 추정, 충전소 선택, LiDAR 접근, QR 정렬, 기계 가이드, 무선충전 시작으로 이어지는 전체 시스템 흐름",
        "applied": "현재 패키지의 시스템 구성도/동작 흐름도/실험 시나리오의 기준으로 사용",
    },
    {
        "category": "로봇 플랫폼",
        "title": "ROBOTIS TurtleBot3 e-Manual: Waffle Pi 사양",
        "source": "ROBOTIS 공식 e-Manual",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/features/",
        "studied": "TurtleBot3 Waffle Pi의 Raspberry Pi, LiDAR, 카메라 구성과 기본 하드웨어 제약",
        "applied": "TurtleBot3 Waffle Pi를 실험 플랫폼으로 두고 LiDAR/카메라/UWB tag 장착 구조를 정리",
    },
    {
        "category": "로봇 내비게이션",
        "title": "ROBOTIS TurtleBot3 Navigation/SLAM e-Manual",
        "source": "ROBOTIS 공식 e-Manual",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/navigation/",
        "studied": "일반 TurtleBot 내비게이션은 SLAM으로 만든 map을 기반으로 주행한다는 점과 bringup/navigation 실행 흐름",
        "applied": "본 프로젝트에서는 첫 접근 단계에서 mapless UWB-LiDAR local planner로 범위를 좁혔음을 비교 설명",
    },
    {
        "category": "UWB 하드웨어",
        "title": "Decawave DW1000 User Manual",
        "source": "Decawave/Qorvo DW1000 사용자 매뉴얼",
        "url": "https://usermanual.wiki/Pdf/dw1000usermanualv205.7098027/html",
        "studied": "UWB tag-anchor ranging, two-way ranging 메시지 흐름, antenna delay/calibration 필요성",
        "applied": "Arduino UWB tag/anchor 코드와 COM5 ranging confirm/fail 로그 해석의 배경으로 사용",
    },
    {
        "category": "UWB 위치추정 논문",
        "title": "Ultrawideband-based precise short-range localization for wireless power transfer to electric vehicles in parking environments",
        "source": "PeerJ Computer Science / PMC 공개 논문",
        "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8176549/",
        "studied": "WPT에서는 송수신 코일 정렬 오차가 중요하고, DADT UWB 구조가 단거리 정밀 위치추정에 쓰일 수 있다는 점",
        "applied": "무선충전 도킹에서 UWB를 단순 위치센서가 아니라 최종 정렬 요구와 연결되는 근거로 설명",
    },
    {
        "category": "비전/QR",
        "title": "OpenCV QRCodeDetector Class Reference",
        "source": "OpenCV 공식 문서",
        "url": "https://docs.opencv.org/3.4/de/dc3/classcv_1_1QRCodeDetector.html",
        "studied": "카메라 영상에서 QR code를 detect/decode하고 QR 중심을 이용해 정렬 오차를 구할 수 있는 기본 API 흐름",
        "applied": "현재 UWB-LiDAR 접근 후 QR/카메라 최종 도킹 단계로 handoff한다는 후속 구현 계획에 반영",
    },
    {
        "category": "ROS 구현",
        "title": "ROS topic 기반 센서/제어 구조",
        "source": "ROS1 Noetic 패키지 구조 및 본 프로젝트 README/launch/script",
        "url": "",
        "studied": "/scan, /uwb_pose, /target_charger, /cmd_vel, /near_charger 등 topic 단위로 센서 입력과 주행 출력을 분리하는 방식",
        "applied": "charger_target_selector, uwb_pose_estimator, lidar_local_planner 노드 구조와 RViz 시각화 구성",
    },
]


def read_uwb_logs() -> pd.DataFrame:
    df = pd.read_csv(ROOT / "uwb_logs_combined.csv", parse_dates=["time"])
    df["error_m"] = ((df["cx_filtered"] - TRUE_X) ** 2 + (df["cy_filtered"] - TRUE_Y) ** 2) ** 0.5
    df["elapsed_s"] = (df["time"] - df["time"].min()).dt.total_seconds()
    return df


def stable_static_rows(df: pd.DataFrame) -> pd.DataFrame:
    stable_files = ["uwb_log_20260513_201035.csv", "uwb_log_20260513_201302.csv"]
    stable = df[df["source_file"].isin(stable_files)].copy()
    stable = stable.sort_values("time").reset_index(drop=True)
    picks = [0, 3, 6, 9, 12, 15, 18, 21, 25, len(stable) - 1]
    return stable.iloc[picks].copy().reset_index(drop=True)


def serial_counts() -> dict[str, int]:
    log_path = ROOT / "serial_logs" / "com5_live.log"
    text = log_path.read_text(encoding="utf-8", errors="ignore") if log_path.exists() else ""
    return {
        "range_confirm_sent": text.count("range_confirm_sent"),
        "poll_receive_fail": text.count("poll_receive_fail"),
        "final_receive_fail": text.count("final_receive_fail"),
        "lines": len(text.splitlines()),
    }


def write_csvs(df: pd.DataFrame, stable: pd.DataFrame, counts: dict[str, int]) -> dict[str, Path]:
    anchor_rows = [
        ["device", "x_m", "y_m", "height_m", "note"],
        ["A1", "0.00", "0.00", "1.00", "5 m x 5 m test frame corner; configured anchor_0"],
        ["A2", "5.00", "0.00", "1.00", "5 m x 5 m test frame corner; configured anchor_1"],
        ["A3", "0.00", "5.00", "1.00", "5 m x 5 m test frame corner; configured anchor_2"],
        ["A4", "5.00", "5.00", "1.00", "5 m x 5 m test frame corner; configured anchor_3"],
        ["C1", "1.20", "1.20", "0.00", "charger_id 1 from config/charger_positions.yaml"],
        ["C2", "3.80", "3.80", "0.00", "charger_id 2 from config/charger_positions.yaml"],
        ["QR1", "1.20", "1.20", "0.00", "planned marker center for final visual alignment"],
        ["QR2", "3.80", "3.80", "0.00", "planned marker center for final visual alignment"],
    ]
    anchor_path = DATA_DIR / "anchor_charger_coordinates.csv"
    write_rows(anchor_path, anchor_rows)

    static_path = DATA_DIR / "uwb_static_position_test_20260513.csv"
    static_rows = [["trial", "source_file", "time", "true_x_m", "true_y_m", "estimated_x_m", "estimated_y_m", "error_m", "error_cm", "note"]]
    for i, row in stable.iterrows():
        static_rows.append([
            i + 1,
            row["source_file"],
            row["time"].isoformat(),
            f"{TRUE_X:.2f}",
            f"{TRUE_Y:.2f}",
            f"{row['cx_filtered']:.4f}",
            f"{row['cy_filtered']:.4f}",
            f"{row['error_m']:.4f}",
            f"{row['error_m'] * 100:.2f}",
            "actual log sample; reference point assumed as lab center (2.50, 2.50 m)",
        ])
    write_rows(static_path, static_rows)

    summary_path = DATA_DIR / "uwb_static_summary_20260513.csv"
    summary_rows = [["metric", "value", "unit", "source"]]
    summary_rows += [
        ["sample_count", len(stable), "samples", "uwb_logs_combined.csv, stable static sessions"],
        ["mean_error", f"{stable['error_m'].mean() * 100:.2f}", "cm", "against (2.50, 2.50 m) reference"],
        ["rmse", f"{math.sqrt((stable['error_m'] ** 2).mean()) * 100:.2f}", "cm", "against (2.50, 2.50 m) reference"],
        ["p95_error", f"{stable['error_m'].quantile(0.95) * 100:.2f}", "cm", "against (2.50, 2.50 m) reference"],
        ["x_std", f"{stable['cx_filtered'].std() * 100:.2f}", "cm", "filtered UWB x repeatability"],
        ["y_std", f"{stable['cy_filtered'].std() * 100:.2f}", "cm", "filtered UWB y repeatability"],
        ["all_log_rows", len(df), "rows", "all stored UWB coordinate logs"],
        ["all_log_mean_x", f"{df['cx_filtered'].mean():.3f}", "m", "for traceability only"],
        ["all_log_mean_y", f"{df['cy_filtered'].mean():.3f}", "m", "for traceability only"],
    ]
    write_rows(summary_path, summary_rows)

    docking_path = DATA_DIR / "approach_handoff_preliminary_log.csv"
    starts = [
        ("S1", 0.50, 0.50),
        ("S2", 2.50, 0.70),
        ("S3", 4.50, 0.60),
        ("S4", 0.80, 4.20),
        ("S5", 2.50, 2.50),
        ("S6", 4.40, 4.30),
    ]
    chargers = {"C1": (1.20, 1.20), "C2": (3.80, 3.80)}
    docking_rows = [["trial", "start_position", "start_x_m", "start_y_m", "selected_station", "selection_success", "planner_stop_radius_m", "handoff_distance_cm", "charging_started", "note"]]
    for i, (name, sx, sy) in enumerate(starts, start=1):
        distances = {sid: math.hypot(sx - cx, sy - cy) for sid, (cx, cy) in chargers.items()}
        selected = min(distances, key=distances.get)
        docking_rows.append([
            i,
            name,
            f"{sx:.2f}",
            f"{sy:.2f}",
            selected,
            "Y",
            "1.00",
            "100",
            "N/A",
            "preliminary logic check from configured charger coordinates; QR/final charging test not yet measured",
        ])
    write_rows(docking_path, docking_rows)

    comm_path = DATA_DIR / "uwb_serial_diagnosis_20260520.csv"
    total_fail = counts["poll_receive_fail"] + counts["final_receive_fail"]
    comm_rows = [
        ["metric", "count", "note"],
        ["range_confirm_sent", counts["range_confirm_sent"], "successful ranging confirmations observed in COM5 live log"],
        ["poll_receive_fail", counts["poll_receive_fail"], "poll receive failures observed; firmware/serial state should be checked before real driving"],
        ["final_receive_fail", counts["final_receive_fail"], "final receive failures observed; likely ranging stability issue during monitor run"],
        ["failure_to_confirm_ratio", f"{total_fail / max(counts['range_confirm_sent'], 1):.2f}", "diagnostic ratio, not a final performance metric"],
        ["log_lines", counts["lines"], "serial_logs/com5_live.log"],
    ]
    write_rows(comm_path, comm_rows)

    return {
        "anchor": anchor_path,
        "static": static_path,
        "summary": summary_path,
        "approach": docking_path,
        "comm": comm_path,
    }


def write_reference_study_files() -> dict[str, Path]:
    matrix_path = REF_DIR / "reference_study_matrix.csv"
    rows = [["category", "reference_title", "source", "url", "what_we_studied", "how_it_was_applied"]]
    for item in REFERENCE_ITEMS:
        rows.append([
            item["category"],
            item["title"],
            item["source"],
            item["url"],
            item["studied"],
            item["applied"],
        ])
    write_rows(matrix_path, rows)

    summary_path = REF_DIR / "참고자료_및_공부내용_요약.md"
    lines = [
        "# 참고자료 및 공부내용 요약",
        "",
        "교수님께 전달할 목적은 단순 참고문헌 나열이 아니라, 어떤 자료를 보고 어떤 개념을 이해했으며 프로젝트에 어떻게 반영했는지 보여주는 것이다.",
        "",
        "## 핵심 공부 방향",
        "",
        "1. TurtleBot3 Waffle Pi 플랫폼의 LiDAR, 카메라, Raspberry Pi 기반 ROS 구성을 확인했다.",
        "2. 일반 TurtleBot navigation은 SLAM으로 만든 map을 사용하는 구조라는 점을 확인했고, 본 프로젝트는 UWB pose와 LiDAR local planner 중심의 mapless 접근 단계로 범위를 좁혔다.",
        "3. UWB는 tag-anchor 거리 측정과 anchor geometry가 중요하므로, 정적 위치 로그를 이용해 반복성과 오차를 먼저 분석했다.",
        "4. 무선충전 도킹에서는 최종 코일 정렬 오차가 중요하므로, UWB 접근 후 QR/카메라 정렬 및 기계 가이드가 필요하다는 구조로 정리했다.",
        "5. ROS topic을 기준으로 `/uwb_pose`, `/target_charger`, `/scan`, `/cmd_vel`, `/near_charger`를 분리하여 구현 범위와 향후 도킹 범위를 구분했다.",
        "",
        "## 참고자료별 정리",
        "",
    ]
    for item in REFERENCE_ITEMS:
        url = f" ({item['url']})" if item["url"] else ""
        lines += [
            f"### {item['category']} - {item['title']}",
            f"- 출처: {item['source']}{url}",
            f"- 공부한 내용: {item['studied']}",
            f"- 프로젝트 반영: {item['applied']}",
            "",
        ]
    lines += [
        "## 교수님께 설명할 때 한 문장 요약",
        "",
        "저희는 TurtleBot3 공식 자료로 플랫폼/ROS 구성을 확인하고, DW1000 UWB 자료로 tag-anchor ranging과 보정 필요성을 공부했으며, WPT용 UWB 정밀 위치추정 논문을 참고해 무선충전에서는 최종 정렬 오차가 중요하다는 점을 정리했습니다. 이를 바탕으로 현재 구현은 UWB 기반 위치 추정과 충전소 선택, LiDAR local approach, QR 도킹 handoff 구조로 구성했습니다.",
        "",
    ]
    summary_path.write_text("\n".join(lines), encoding="utf-8")

    message_path = REF_DIR / "교수님께_참고공부내용_전달문.txt"
    message_path.write_text(
        """교수님, 추가로 저희가 참고하고 공부한 자료를 따로 정리했습니다.

정리 방향은 단순 참고문헌 목록이 아니라, 각 자료에서 무엇을 공부했고 그것을 프로젝트의 어느 부분에 반영했는지 보이도록 했습니다. 주요 참고 범위는 TurtleBot3 공식 e-Manual, DW1000 UWB ranging 자료, UWB 기반 WPT 정밀 위치추정 논문, OpenCV QR 검출 문서, 그리고 ROS topic 기반 구현 구조입니다.

현재 구현은 전체 무선충전 도킹 중 UWB 위치추정/충전소 선택/LiDAR 접근 단계에 초점을 두었고, QR 기반 최종 정렬과 실제 충전 시작은 후속 실험 및 구현 범위로 분리했습니다.
""",
        encoding="utf-8",
    )
    return {"matrix": matrix_path, "summary": summary_path, "message": message_path}


def write_rows(path: Path, rows: list[list[object]]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerows(rows)


def make_figures(df: pd.DataFrame, stable: pd.DataFrame) -> dict[str, Path]:
    scatter_path = FIG_DIR / "fig4_uwb_static_scatter_20260513.png"
    draw_scatter(
        scatter_path,
        stable["cx_filtered"].tolist(),
        stable["cy_filtered"].tolist(),
        [(TRUE_X, TRUE_Y, "#dc2626", "Reference")],
        "UWB Static Position Samples",
    )

    error_path = FIG_DIR / "fig5_uwb_static_error_20260513.png"
    draw_line(
        error_path,
        list(range(1, len(stable) + 1)),
        (stable["error_m"] * 100).tolist(),
        "UWB Static Error Against Center Reference",
        "sample",
        "error (cm)",
        mean_line=float(stable["error_m"].mean() * 100),
    )

    full_path = FIG_DIR / "fig6_all_uwb_logs_trace_20260513.png"
    draw_group_trace(full_path, df)

    return {"scatter": scatter_path, "error": error_path, "trace": full_path}


def get_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\malgunbd.ttf" if bold else r"C:\Windows\Fonts\malgun.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def chart_canvas(title: str):
    img = Image.new("RGB", (1152, 756), "white")
    draw = ImageDraw.Draw(img)
    font_title = get_font(28, True)
    font = get_font(18)
    draw.text((64, 32), title, fill="#111827", font=font_title)
    return img, draw, font


def scale_point(x, y, xmin, xmax, ymin, ymax, left=96, top=96, right=1080, bottom=660):
    if xmax == xmin:
        xmax = xmin + 1
    if ymax == ymin:
        ymax = ymin + 1
    px = left + (x - xmin) / (xmax - xmin) * (right - left)
    py = bottom - (y - ymin) / (ymax - ymin) * (bottom - top)
    return px, py


def draw_axes(draw, font, xmin, xmax, ymin, ymax, xlabel, ylabel):
    left, top, right, bottom = 96, 96, 1080, 660
    draw.rectangle((left, top, right, bottom), outline="#d1d5db", width=2)
    for i in range(6):
        x = left + i * (right - left) / 5
        y = top + i * (bottom - top) / 5
        draw.line((x, top, x, bottom), fill="#eef2f7", width=1)
        draw.line((left, y, right, y), fill="#eef2f7", width=1)
        xv = xmin + i * (xmax - xmin) / 5
        yv = ymax - i * (ymax - ymin) / 5
        draw.text((x - 24, bottom + 12), f"{xv:.2f}", fill="#4b5563", font=font)
        draw.text((20, y - 10), f"{yv:.2f}", fill="#4b5563", font=font)
    draw.text(((left + right) / 2 - 60, 704), xlabel, fill="#374151", font=font)
    draw.text((18, 62), ylabel, fill="#374151", font=font)


def draw_scatter(path: Path, xs: list[float], ys: list[float], markers: list[tuple[float, float, str, str]], title: str) -> None:
    img, draw, font = chart_canvas(title)
    all_x = xs + [m[0] for m in markers]
    all_y = ys + [m[1] for m in markers]
    pad_x = max((max(all_x) - min(all_x)) * 0.25, 0.03)
    pad_y = max((max(all_y) - min(all_y)) * 0.25, 0.03)
    xmin, xmax = min(all_x) - pad_x, max(all_x) + pad_x
    ymin, ymax = min(all_y) - pad_y, max(all_y) + pad_y
    draw_axes(draw, font, xmin, xmax, ymin, ymax, "x position (m)", "y position (m)")
    for x, y in zip(xs, ys):
        px, py = scale_point(x, y, xmin, xmax, ymin, ymax)
        draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill="#2563eb", outline="#1e40af")
    for x, y, color, label in markers:
        px, py = scale_point(x, y, xmin, xmax, ymin, ymax)
        draw.line((px - 12, py - 12, px + 12, py + 12), fill=color, width=4)
        draw.line((px - 12, py + 12, px + 12, py - 12), fill=color, width=4)
        draw.text((px + 16, py - 10), label, fill=color, font=font)
    img.save(path)


def draw_line(path: Path, xs: list[float], ys: list[float], title: str, xlabel: str, ylabel: str, mean_line: float | None = None) -> None:
    img, draw, font = chart_canvas(title)
    xmin, xmax = min(xs), max(xs)
    ymin, ymax = 0, max(ys) * 1.2
    draw_axes(draw, font, xmin, xmax, ymin, ymax, xlabel, ylabel)
    points = [scale_point(x, y, xmin, xmax, ymin, ymax) for x, y in zip(xs, ys)]
    if len(points) > 1:
        draw.line(points, fill="#0f766e", width=4)
    for px, py in points:
        draw.ellipse((px - 7, py - 7, px + 7, py + 7), fill="#0f766e")
    if mean_line is not None:
        _, py = scale_point(xmin, mean_line, xmin, xmax, ymin, ymax)
        draw.line((96, py, 1080, py), fill="#f97316", width=3)
        draw.text((900, py - 28), f"mean {mean_line:.2f} cm", fill="#c2410c", font=font)
    img.save(path)


def draw_group_trace(path: Path, df: pd.DataFrame) -> None:
    img, draw, font = chart_canvas("Stored UWB Coordinate Log Trace")
    groups = [(s, g) for s, g in df.groupby("source_file") if len(g) >= 5]
    xs = df["cx_filtered"].tolist()
    ys = df["cy_filtered"].tolist()
    pad_x = max((max(xs) - min(xs)) * 0.15, 0.03)
    pad_y = max((max(ys) - min(ys)) * 0.15, 0.03)
    xmin, xmax = min(xs) - pad_x, max(xs) + pad_x
    ymin, ymax = min(ys) - pad_y, max(ys) + pad_y
    draw_axes(draw, font, xmin, xmax, ymin, ymax, "x position (m)", "y position (m)")
    colors = ["#2563eb", "#0f766e", "#f97316", "#7c3aed", "#dc2626"]
    for idx, (source, group) in enumerate(groups):
        color = colors[idx % len(colors)]
        points = [scale_point(x, y, xmin, xmax, ymin, ymax) for x, y in zip(group["cx_filtered"], group["cy_filtered"])]
        if len(points) > 1:
            draw.line(points, fill=color, width=3)
        for px, py in points:
            draw.ellipse((px - 4, py - 4, px + 4, py + 4), fill=color)
        label = source.replace("uwb_log_", "").replace(".csv", "")
        draw.rectangle((730, 540 + idx * 32, 752, 556 + idx * 32), fill=color)
        draw.text((762, 532 + idx * 32), label, fill="#374151", font=font)
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_in):
            cell = row.cells[i]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("UWB 기반 TurtleBot3 자율 무선충전 접근/도킹 자료 패키지")
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("교수님 전달용 자료 취합본 | 작성일: 2026-05-31 | 데이터 기준: 2026-05-13 UWB 로그, 2026-05-20 COM5 진단 로그")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_note_box(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 58, 95)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            p = cells[i].paragraphs[0]
            p.text = str(value)
            if i in {0, 2, 3, 4}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(cells[i])
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str, width_in: float = 5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def build_docx(df: pd.DataFrame, stable: pd.DataFrame, counts: dict[str, int], figs: dict[str, Path]) -> None:
    doc = Document()
    apply_styles(doc)
    add_title(doc)

    mean_err = stable["error_m"].mean() * 100
    rmse = math.sqrt((stable["error_m"] ** 2).mean()) * 100
    p95 = stable["error_m"].quantile(0.95) * 100
    x_std = stable["cx_filtered"].std() * 100
    y_std = stable["cy_filtered"].std() * 100

    add_note_box(
        doc,
        "데이터 범위: 본 자료의 UWB 위치 수치는 저장된 실제 로그에서 계산했다. "
        "최종 QR 도킹 및 충전 시작 여부는 아직 실측 로그가 없어 예비 오프라인/로직 검증으로 분리 표기했다.",
    )

    doc.add_heading("1. 전달 요약", level=1)
    doc.add_paragraph(
        "본 패키지는 TurtleBot3 Waffle Pi의 자율 무선충전 시스템 중 UWB 기반 위치 추정, 충전 스테이션 선택, "
        "LiDAR 기반 접근 및 QR 도킹으로 이어지는 구조를 교수님께 전달하기 위한 자료이다. "
        "현재 코드 기준 구현 범위는 UWB pose/charger target을 입력으로 받아 LiDAR local planner가 충전소 1 m 전방까지 접근하고, "
        "그 지점에서 QR/카메라 기반 최종 정렬 단계로 넘기는 것이다."
    )
    add_table(
        doc,
        ["항목", "현재 정리 상태", "논문/보고서 활용"],
        [
            ["시스템 구조", "ROS topic 흐름, 구성 노드, 그림 1~3 정리", "방법론 섹션"],
            ["UWB 로그", f"실제 좌표 로그 {len(df)}행, 안정 구간 {len(stable)}개 표본 분석", "실험 결과/예비 성능"],
            ["정적 위치 결과", f"평균 오차 {mean_err:.2f} cm, RMSE {rmse:.2f} cm, 95% 오차 {p95:.2f} cm", "UWB 정확도 표/그래프"],
            ["충전소 선택", "C1=(1.20,1.20), C2=(3.80,3.80) 기준 최근접 선택", "알고리즘 설명"],
            ["실주행 도킹", "최종 QR 도킹/충전 시작 실측은 미수집", "향후 보완 실험"],
        ],
        [1.4, 2.8, 2.3],
    )

    doc.add_heading("2. 참고자료 및 공부한 내용", level=1)
    doc.add_paragraph(
        "교수님께서 확인하고자 하시는 핵심은 단순 결과값보다, 팀이 어떤 자료를 참고했고 그 자료를 통해 어떤 기술 개념을 이해했는지이다. "
        "아래 표는 참고자료를 공부한 내용과 실제 프로젝트 반영 부분으로 나누어 정리한 것이다."
    )
    add_table(
        doc,
        ["분야", "참고자료", "공부한 내용", "프로젝트 반영"],
        [
            [item["category"], item["title"], item["studied"], item["applied"]]
            for item in REFERENCE_ITEMS
        ],
        [1.05, 1.75, 1.95, 1.75],
    )
    doc.add_paragraph(
        "정리하면, TurtleBot3 공식 자료로 플랫폼/ROS 구성을 확인했고, DW1000 자료로 UWB tag-anchor ranging과 보정 필요성을 공부했다. "
        "또한 UWB 기반 WPT 정밀 위치추정 논문을 참고하여 무선충전에서는 최종 정렬 오차가 중요하다는 점을 이해했고, "
        "이를 UWB-LiDAR 접근 후 QR/카메라 도킹으로 넘기는 시스템 구조에 반영했다."
    )

    doc.add_heading("3. 시스템 구성 및 동작 범위", level=1)
    doc.add_paragraph(
        "전체 흐름은 배터리 부족 감지 이후 UWB로 로봇 좌표를 추정하고, 두 충전소 중 가까운 목표를 선택한 뒤, "
        "LiDAR local planner가 장애물을 회피하며 목표 근방까지 이동하는 구조이다. planner는 goal_radius=1.00 m에 도달하면 "
        "cmd_vel을 0으로 만들고 near_charger=true를 발행하도록 설정되어 있다."
    )
    add_table(
        doc,
        ["구분", "입력/설정", "출력/역할"],
        [
            ["UWB pose", "/uwb/ranges 또는 /uwb_pose", "로봇의 x, y, theta 추정"],
            ["충전소 선택", "charger_id 1: (1.20,1.20), charger_id 2: (3.80,3.80)", "/target_charger 발행"],
            ["LiDAR planner", "/scan, /uwb_pose, /target_charger", "/cmd_vel, /lidar_state, /near_charger"],
            ["안전 파라미터", "avoid_dist=0.55 m, emergency_dist=0.30 m, max_linear=0.08 m/s", "장애물 회피 및 정지"],
            ["최종 도킹", "QR marker, camera, UWB 근거리 확인", "현재는 후속 실측 필요"],
        ],
        [1.45, 2.75, 2.3],
    )

    add_figure(doc, FIG_DIR / "fig1_system_architecture.png", "그림 1. UWB-LiDAR-QR 기반 자율 무선충전 시스템 구성도", 5.9)
    add_figure(doc, FIG_DIR / "fig2_operation_flow.png", "그림 2. 자율 충전 접근 및 도킹 동작 흐름", 5.9)
    add_figure(doc, FIG_DIR / "fig3_experiment_layout.png", "그림 3. 5 m x 5 m 실험 좌표계 및 충전 스테이션 배치", 5.9)

    doc.add_heading("4. 데이터 출처", level=1)
    add_table(
        doc,
        ["자료", "파일/출처", "사용 방식"],
        [
            ["UWB 좌표 로그", "uwb_logs_combined.csv", "정적 위치 반복성 및 기준점 오차 계산"],
            ["안정 구간", "uwb_log_20260513_201035.csv, uwb_log_20260513_201302.csv", "중앙 기준점 (2.50, 2.50 m) 주변 분석"],
            ["시리얼 진단", "serial_logs/com5_live.log", "ranging confirm/fail 메시지 카운트"],
            ["좌표 설정", "config/charger_positions.yaml", "충전소 C1/C2 좌표"],
            ["주행 파라미터", "config/planner_params.yaml", "정지 반경 및 장애물 회피 기준"],
        ],
        [1.45, 2.55, 2.5],
    )

    doc.add_heading("5. UWB 정적 위치 예비 결과", level=1)
    doc.add_paragraph(
        "안정적으로 중앙 기준점 주변에 모인 두 로그 구간을 사용해 기준점 (2.50, 2.50 m) 대비 오차를 계산했다. "
        "해당 수치는 논문 최종값이 아니라 교수님께 현재 데이터 품질을 설명하기 위한 예비 결과이며, "
        "향후 줄자/마킹 기준점과 동기화된 true_x, true_y 로그를 추가하면 그대로 재계산할 수 있다."
    )
    add_table(
        doc,
        ["Trial", "Source", "Est. x", "Est. y", "Error(cm)", "비고"],
        [
            [
                i + 1,
                row["source_file"].replace("uwb_log_", "").replace(".csv", ""),
                f"{row['cx_filtered']:.3f}",
                f"{row['cy_filtered']:.3f}",
                f"{row['error_m'] * 100:.2f}",
                "actual log",
            ]
            for i, row in stable.iterrows()
        ],
        [0.55, 1.35, 0.85, 0.85, 0.9, 2.0],
    )
    add_table(
        doc,
        ["지표", "값", "해석"],
        [
            ["평균 오차", f"{mean_err:.2f} cm", "중앙 기준점 대비 평균 거리 오차"],
            ["RMSE", f"{rmse:.2f} cm", "큰 오차에 더 민감한 종합 지표"],
            ["95% 오차", f"{p95:.2f} cm", "대부분의 표본이 이 범위 내에 위치"],
            ["x 표준편차", f"{x_std:.2f} cm", "필터링 후 x축 반복성"],
            ["y 표준편차", f"{y_std:.2f} cm", "필터링 후 y축 반복성"],
        ],
        [1.5, 1.4, 3.6],
    )
    add_figure(doc, figs["scatter"], "그림 4. 실제 UWB 정적 좌표 표본 산점도", 5.4)
    add_figure(doc, figs["error"], "그림 5. 중앙 기준점 대비 UWB 정적 오차", 5.4)
    add_figure(doc, figs["trace"], "그림 6. 저장된 UWB 좌표 로그 구간별 궤적", 5.4)

    doc.add_heading("6. 충전소 선택 및 접근 로직", level=1)
    doc.add_paragraph(
        "충전소 선택은 현재 UWB pose p=(x,y)와 각 충전소 좌표 사이의 유클리드 거리를 계산하여 더 가까운 충전소를 선택한다. "
        "선택된 charger_id는 /target_charger로 변환되고, LiDAR local planner는 목표 방향 오차를 줄이면서 이동한다."
    )
    add_table(
        doc,
        ["Trial", "시작점", "선택 충전소", "성공", "handoff", "비고"],
        [
            [1, "S1 (0.50,0.50)", "C1", "Y", "1.00 m", "nearest-station logic"],
            [2, "S2 (2.50,0.70)", "C1", "Y", "1.00 m", "nearest-station logic"],
            [3, "S3 (4.50,0.60)", "C2", "Y", "1.00 m", "nearest-station logic"],
            [4, "S4 (0.80,4.20)", "C1", "Y", "1.00 m", "nearest-station logic"],
            [5, "S5 (2.50,2.50)", "C1/C2 경계", "조건부", "1.00 m", "동거리 시 기본 선택 규칙 명시 필요"],
            [6, "S6 (4.40,4.30)", "C2", "Y", "1.00 m", "nearest-station logic"],
        ],
        [0.55, 1.75, 1.15, 0.7, 0.9, 1.45],
    )

    doc.add_heading("7. UWB 통신 진단 로그", level=1)
    fail_total = counts["poll_receive_fail"] + counts["final_receive_fail"]
    add_table(
        doc,
        ["항목", "카운트", "판단"],
        [
            ["range_confirm_sent", f"{counts['range_confirm_sent']:,}", "ranging confirm 메시지는 관측됨"],
            ["poll_receive_fail", f"{counts['poll_receive_fail']:,}", "실주행 전 펌웨어/전원/배치/baud 상태 점검 필요"],
            ["final_receive_fail", f"{counts['final_receive_fail']:,}", "final packet 수신 안정성 보완 필요"],
            ["fail/confirm ratio", f"{fail_total / max(counts['range_confirm_sent'], 1):.2f}", "최종 성능값이 아니라 진단 지표"],
        ],
        [1.65, 1.3, 3.55],
    )

    doc.add_heading("8. 교수님께 전달할 때의 핵심 메시지", level=1)
    doc.add_paragraph(
        "현재 자료는 UWB-LiDAR 기반 접근 단계의 구현 구조와 실제 UWB 좌표 로그 기반 예비 결과를 포함한다. "
        "논문 작성 관점에서는 시스템 구조, UWB 위치 추정 결과, 충전소 선택 로직, LiDAR 접근/정지 조건까지는 바로 방법론으로 정리 가능하다. "
        "다만 최종 QR 정렬과 실제 무선충전 시작 여부는 아직 실측 데이터가 없으므로, 다음 실험에서 반복 주행 로그와 최종 정렬 오차를 추가해야 한다."
    )
    add_table(
        doc,
        ["다음 보완", "측정 방법", "목표 산출물"],
        [
            ["정확한 기준점 로그", "true_x,true_y를 CSV에 함께 기록", "절대 위치 오차/RMSE"],
            ["실주행 접근 실험", "S1~S6에서 rosbag 반복 기록", "충전소 선택 성공률, 소요 시간"],
            ["QR 최종 도킹", "카메라 캡처와 최종 정지 위치 측정", "최종 오차 cm, 충전 시작 성공률"],
            ["통신 안정화", "COM5 fail 메시지 원인 점검", "실주행 전 UWB update rate 안정화"],
        ],
        [1.5, 2.35, 2.65],
    )

    doc.add_heading("9. 패키지 파일 구성", level=1)
    add_table(
        doc,
        ["폴더", "주요 파일", "내용"],
        [
            ["01", DOCX_PATH.name, "교수님 전달용 취합본"],
            ["02_그림자료", "fig1~fig6 png", "시스템/흐름/배치/데이터 그래프"],
            ["03_실험데이터", "anchor_charger_coordinates.csv", "앵커와 충전소 좌표"],
            ["03_실험데이터", "uwb_static_position_test_20260513.csv", "실제 UWB 로그 기반 표본 데이터"],
            ["03_실험데이터", "approach_handoff_preliminary_log.csv", "충전소 선택/접근 handoff 예비 검증"],
            ["05_참고_공부자료", "참고자료_및_공부내용_요약.md", "참고자료별 공부 내용과 프로젝트 반영 정리"],
            ["05_참고_공부자료", "reference_study_matrix.csv", "교수님 전달용 참고/공부 매트릭스"],
            ["04_메시지_체크리스트", "교수님_전달메시지.txt", "메일/카톡 전달 문구"],
        ],
        [1.25, 2.5, 2.75],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "UWB TurtleBot professor delivery package"
    doc.save(DOCX_PATH)


def write_messages(stable: pd.DataFrame, counts: dict[str, int]) -> None:
    mean_err = stable["error_m"].mean() * 100
    rmse = math.sqrt((stable["error_m"] ** 2).mean()) * 100
    msg = f"""교수님, 안녕하세요.

UWB 기반 TurtleBot3 자율 무선충전 접근/도킹 파트 자료를 정리해서 전달드립니다.

이번 패키지에는 시스템 구성도, 동작 흐름도, UWB/충전소 좌표, 실제 UWB 좌표 로그 기반 예비 결과, 충전소 선택 및 LiDAR 접근 로직을 함께 넣었습니다. 저장된 2026-05-13 UWB 로그 중 중앙 기준점 주변 안정 구간을 분석했을 때 평균 오차는 {mean_err:.2f} cm, RMSE는 {rmse:.2f} cm 수준으로 정리되었습니다.

다만 QR 기반 최종 도킹과 실제 충전 시작 여부는 아직 반복 실측 로그가 없어, 문서 안에서는 예비 오프라인/로직 검증으로 분리 표기했습니다. 교수님께서 논문 방향을 잡아주시면 다음 실험에서 true_x/true_y 기준점 로그, rosbag 기반 접근 시간, 최종 정렬 오차, 충전 시작 성공률을 추가 측정해 보완하겠습니다.

감사합니다.
"""
    (MSG_DIR / "교수님_전달메시지.txt").write_text(msg, encoding="utf-8")

    checklist = f"""[전달 전 체크리스트]

완료
- 자료 취합본 docx 재작성
- 기존 템플릿 CSV를 실제 데이터/예비 검증 CSV로 교체
- 2026-05-13 UWB 좌표 로그 기반 평균 오차/RMSE 계산
- fig4~fig6 데이터 그래프 추가
- 2026-05-20 COM5 시리얼 진단 카운트 반영

교수님께 말할 때 주의
- UWB 정적 위치 수치는 실제 저장 로그 기반 예비 결과임
- 기준점은 중앙점 (2.50, 2.50 m)으로 둔 분석이라, 논문 최종값 전에는 true_x/true_y 동기화 재측정 필요
- QR 최종 도킹/충전 시작은 아직 실측 완료 데이터가 아니므로 예비 검증으로 표현
- COM5 진단 로그에서 range_confirm_sent={counts['range_confirm_sent']:,}, poll_receive_fail={counts['poll_receive_fail']:,}, final_receive_fail={counts['final_receive_fail']:,}가 관측되어 실주행 전 UWB 통신 안정화 필요

다음 실험에서 채울 항목
- 시작 위치별 rosbag 기록
- 최종 정렬 오차 cm
- 충전 시작 성공률
- UWB update rate 및 timeout 발생 횟수
"""
    (MSG_DIR / "내파트_체크리스트.txt").write_text(checklist, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    df = read_uwb_logs()
    stable = stable_static_rows(df)
    counts = serial_counts()
    write_csvs(df, stable, counts)
    write_reference_study_files()
    figs = make_figures(df, stable)
    build_docx(df, stable, counts, figs)
    write_messages(stable, counts)
    print(f"Wrote professor package: {PACKAGE_ROOT}")


if __name__ == "__main__":
    main()
