from __future__ import annotations

import csv
import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "UWB_TurtleBot_교수님전달용_논문초안_수식포함"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "01_UWB_TurtleBot_논문초안_수식포함_교수님검토용.docx"
EQUATION_PATH = OUT_DIR / "02_수식_정리표.csv"
REF_PATH = OUT_DIR / "03_참고문헌_정리.csv"
MSG_PATH = OUT_DIR / "04_교수님_전달메시지.txt"
ZIP_PATH = ROOT / "UWB_TurtleBot_교수님전달용_논문초안_수식포함.zip"

SOURCE_FIG_DIRS = [
    ROOT / "UWB_TurtleBot_교수님전달용_참고공부자료_그림추가" / "05_논문그림자료",
    ROOT / "UWB_TurtleBot_교수님전달용_자료패키지" / "UWB_TurtleBot_교수님전달용_자료패키지" / "02_그림자료",
]

TRUE_X = 2.50
TRUE_Y = 2.50


REFERENCES = [
    {
        "id": "[1]",
        "title": "ROBOTIS TurtleBot3 e-Manual: Features / Waffle Pi",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/features/",
        "use": "TurtleBot3 Waffle Pi 플랫폼, LiDAR/카메라/ROS 기반 구성 근거",
    },
    {
        "id": "[2]",
        "title": "ROBOTIS TurtleBot3 e-Manual: Navigation",
        "url": "https://emanual.robotis.com/docs/en/platform/turtlebot3/navigation/",
        "use": "일반 SLAM/map 기반 navigation과 본 연구의 mapless UWB-LiDAR 접근 방식 비교",
    },
    {
        "id": "[3]",
        "title": "Qorvo DW1000 UWB Transceiver",
        "url": "https://www.qorvo.com/products/p/DW1000",
        "use": "UWB ranging, TWR/TDoA RTLS, 약 10 cm급 위치 응용 가능성 근거",
    },
    {
        "id": "[4]",
        "title": "Ultrawideband-based precise short-range localization for wireless power transfer to electric vehicles in parking environments",
        "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC8176549/",
        "use": "무선전력전송에서 단거리 위치추정과 정렬 오차가 중요하다는 연구 배경",
    },
    {
        "id": "[5]",
        "title": "OpenCV QRCodeDetector Class Reference",
        "url": "https://docs.opencv.org/3.4/de/dc3/classcv_1_1QRCodeDetector.html",
        "use": "QR marker 검출과 최종 시각 정렬 단계 설계 근거",
    },
]


EQUATIONS = [
    {
        "no": "(1)",
        "name": "UWB 거리 측정 모델",
        "eq": "z_i(k) = ||p(k) - a_i||_2 + b_i + eps_i(k)",
        "desc": "i번째 anchor a_i와 tag 위치 p 사이의 실제 거리, anchor bias b_i, 측정 잡음 eps_i를 포함한 ranging 모델.",
    },
    {
        "no": "(2)",
        "name": "ToF 기반 거리 환산",
        "eq": "r_i = c * tau_i",
        "desc": "전파 전파시간 tau_i에 광속 c를 곱해 거리로 환산한다. 실제 DW1000 TWR에서는 clock offset/antenna delay 보정이 필요하다.",
    },
    {
        "no": "(3)",
        "name": "비선형 최소제곱 위치추정",
        "eq": "p_hat = argmin_p sum_i w_i (||p - a_i||_2 - (z_i - b_i))^2",
        "desc": "4개 UWB anchor 거리값을 가장 잘 만족하는 2D 위치 p를 추정한다.",
    },
    {
        "no": "(4)",
        "name": "선형화된 trilateration 식",
        "eq": "2(a_i-a_1)^T p = r_1^2 - r_i^2 + ||a_i||^2 - ||a_1||^2,  i=2,...,N",
        "desc": "첫 anchor 방정식을 빼서 x,y에 대한 선형 방정식 Ap=b를 구성한다.",
    },
    {
        "no": "(5)",
        "name": "위치 smoothing",
        "eq": "p_f(k) = alpha p_hat(k) + (1-alpha) p_f(k-1)",
        "desc": "본 프로젝트 config의 smoothing_alpha=0.35를 적용할 수 있는 1차 저역통과 필터.",
    },
    {
        "no": "(6)",
        "name": "충전소 선택",
        "eq": "j* = argmin_j ||p_f - c_j||_2,  j in {1,2}",
        "desc": "현재 UWB 위치에서 가장 가까운 충전 스테이션을 target charger로 선택한다.",
    },
    {
        "no": "(7)",
        "name": "목표 heading 오차",
        "eq": "e_theta = wrap(atan2(c_y - y, c_x - x) - theta)",
        "desc": "로봇 heading과 목표 충전소 방향 사이의 각도 오차.",
    },
    {
        "no": "(8)",
        "name": "LiDAR local planner 속도 명령",
        "eq": "v = v_max I(d_goal > R_goal),   omega = sat(k_theta e_theta, -omega_max, omega_max)",
        "desc": "goal_radius 밖에서는 저속 전진하고 heading 오차에 비례해 회전한다.",
    },
    {
        "no": "(9)",
        "name": "LiDAR sector 안전거리",
        "eq": "d_S = min_{phi in S} r(phi)",
        "desc": "front/front-left/front-right 등 sector별 최소 LaserScan 거리로 장애물 회피 상태를 결정한다.",
    },
    {
        "no": "(10)",
        "name": "QR marker 중심 오차",
        "eq": "e_q = [u_q - u_0, v_q - v_0]^T,   q_c = (1/4) sum_m q_m",
        "desc": "QR 네 꼭짓점 평균으로 marker 중심을 구하고 영상 중심과의 픽셀 오차를 최종 정렬에 사용한다.",
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for name, size, color, before, after in [
        ("Heading 1", 14, "1F4D78", 12, 5),
        ("Heading 2", 11.5, "2E74B5", 8, 4),
        ("Heading 3", 10.5, "1F4D78", 6, 3),
    ]:
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell)
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.text = value
            for run in p.runs:
                run.font.size = Pt(font_size)
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_equation(doc: Document, no: str, eq: str, desc: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{no}   {eq}")
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    r.bold = True
    p = doc.add_paragraph(desc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].font.size = Pt(8.5)
        cap.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def copy_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figure_names = [
        "fig01_system_architecture.png",
        "fig02_operation_flow.png",
        "fig03_experiment_layout.png",
        "fig04_reference_to_system_map.png",
        "fig05_uwb_anchor_geometry.png",
        "fig06_ros_topic_pipeline.png",
        "fig07_uwb_static_scatter.png",
        "fig08_uwb_static_error.png",
        "fig09_all_uwb_log_trace.png",
    ]
    copied = {}
    for name in figure_names:
        for source_dir in SOURCE_FIG_DIRS:
            source = source_dir / name
            if source.exists():
                dest = FIG_DIR / name
                shutil.copy2(source, dest)
                copied[name] = dest
                break
    fallback_map = {
        "fig01_system_architecture.png": "fig1_system_architecture.png",
        "fig02_operation_flow.png": "fig2_operation_flow.png",
        "fig03_experiment_layout.png": "fig3_experiment_layout.png",
        "fig07_uwb_static_scatter.png": "fig4_uwb_static_scatter_20260513.png",
        "fig08_uwb_static_error.png": "fig5_uwb_static_error_20260513.png",
        "fig09_all_uwb_log_trace.png": "fig6_all_uwb_logs_trace_20260513.png",
    }
    for dest_name, source_name in fallback_map.items():
        if dest_name in copied:
            continue
        for source_dir in SOURCE_FIG_DIRS:
            source = source_dir / source_name
            if source.exists():
                dest = FIG_DIR / dest_name
                shutil.copy2(source, dest)
                copied[dest_name] = dest
                break
    return copied


def compute_metrics() -> dict[str, float]:
    df = pd.read_csv(ROOT / "uwb_logs_combined.csv", parse_dates=["time"])
    df["error_m"] = ((df["cx_filtered"] - TRUE_X) ** 2 + (df["cy_filtered"] - TRUE_Y) ** 2) ** 0.5
    stable = df[df["source_file"].isin(["uwb_log_20260513_201035.csv", "uwb_log_20260513_201302.csv"])].copy()
    stable = stable.sort_values("time").reset_index(drop=True)
    picks = [0, 3, 6, 9, 12, 15, 18, 21, 25, len(stable) - 1]
    sample = stable.iloc[picks].copy()
    return {
        "n": float(len(sample)),
        "mean_error_cm": float(sample["error_m"].mean() * 100),
        "rmse_cm": float(math.sqrt((sample["error_m"] ** 2).mean()) * 100),
        "p95_cm": float(sample["error_m"].quantile(0.95) * 100),
        "x_std_cm": float(sample["cx_filtered"].std() * 100),
        "y_std_cm": float(sample["cy_filtered"].std() * 100),
    }


def build_docx(figs: dict[str, Path], metrics: dict[str, float]) -> None:
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("UWB 위치추정과 LiDAR Local Planner를 이용한 TurtleBot3 자율 무선충전 접근 시스템")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("교수님 검토용 논문 초안 | 수식/그림/예비 로그 결과 포함 | 2026-05-31")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("초록", level=1)
    doc.add_paragraph(
        "본 연구는 실내 모바일 로봇의 자율 무선충전 시나리오에서 UWB 기반 위치추정, 충전 스테이션 선택, "
        "LiDAR 기반 local approach, QR 기반 최종 정렬 handoff를 결합한 TurtleBot3 Waffle Pi 시스템을 제안한다. "
        "GPS 사용이 어려운 실내 환경에서 UWB anchor-tag ranging은 충전소 접근을 위한 전역 좌표 단서를 제공하고, "
        "LiDAR는 접근 중 국소 장애물 회피를 담당한다. 로봇이 충전소 반경 1 m에 도달하면 정지 후 QR/카메라 정렬 단계로 전환된다. "
        f"저장된 UWB 예비 로그를 중앙 기준점으로 분석한 결과, 평균 오차 {metrics['mean_error_cm']:.2f} cm, "
        f"RMSE {metrics['rmse_cm']:.2f} cm가 관측되었다. 단, 이 값은 true pose가 동기화된 최종 실험값이 아니라 "
        "논문 실험 설계를 위한 예비 로그 분석 결과이다."
    )
    doc.add_paragraph("주요어: UWB, TurtleBot3, wireless charging, LiDAR local planner, QR docking, indoor localization")

    doc.add_heading("1. 서론", level=1)
    doc.add_paragraph(
        "실내 자율 로봇이 장시간 운용되기 위해서는 배터리 부족 시 스스로 충전 위치로 이동하고, 충전 가능한 자세로 정렬하는 기능이 필요하다. "
        "TurtleBot3 Waffle Pi는 ROS 기반 모바일 로봇 플랫폼으로 LiDAR와 카메라를 활용할 수 있으며[1], 일반적인 TurtleBot navigation은 "
        "SLAM으로 생성한 map을 활용한다[2]. 그러나 충전소 접근 단계에서는 전체 map에 의존하지 않고 UWB 좌표와 LiDAR local obstacle avoidance만으로도 "
        "단순 실내 환경에서 충전소 근방까지 접근할 수 있다."
    )
    doc.add_paragraph(
        "본 연구의 기여는 다음과 같다. 첫째, UWB anchor-tag ranging을 이용하여 로봇 pose와 충전소 선택을 결합한 접근 구조를 제시한다. "
        "둘째, LiDAR sector distance를 이용한 local planner로 mapless 접근 단계를 구현한다. 셋째, 무선전력전송(WPT)의 정렬 오차 중요성을 고려하여 "
        "UWB-LiDAR 접근 이후 QR/카메라 기반 최종 정렬 단계로 handoff하는 논문 구조를 제안한다[4][5]."
    )

    doc.add_heading("2. 관련 연구 및 참고자료", level=1)
    add_table(
        doc,
        ["구분", "참고자료", "논문에서 쓰는 역할"],
        [[ref["id"], ref["title"], ref["use"]] for ref in REFERENCES],
        [0.55, 2.8, 3.65],
    )
    add_figure(doc, figs.get("fig04_reference_to_system_map.png", Path()), "Fig. 1. 참고자료와 본 시스템 설계의 연결 관계.", 6.2)

    doc.add_heading("3. 시스템 구성", level=1)
    doc.add_paragraph(
        "시스템은 UWB anchor 4개, TurtleBot3 Waffle Pi에 장착된 UWB tag, LiDAR, 카메라, 두 개의 충전 스테이션으로 구성된다. "
        "UWB는 로봇 위치와 충전소 선택에 사용하고, LiDAR는 local obstacle avoidance 및 충전소 근방 접근에 사용한다. "
        "카메라는 `/near_charger=true` 이후 QR marker를 검출하여 최종 정렬 오차를 줄이는 단계에서 사용된다."
    )
    add_figure(doc, figs.get("fig01_system_architecture.png", Path()), "Fig. 2. UWB-LiDAR-QR 기반 TurtleBot 자율 무선충전 시스템 구성.", 6.2)
    add_figure(doc, figs.get("fig02_operation_flow.png", Path()), "Fig. 3. 배터리 부족부터 QR docking handoff까지의 동작 흐름.", 6.2)
    add_figure(doc, figs.get("fig06_ros_topic_pipeline.png", Path()), "Fig. 4. ROS topic 기반 구현 pipeline.", 6.2)

    doc.add_heading("4. UWB 위치추정 모델", level=1)
    doc.add_paragraph(
        "DW1000 계열 UWB transceiver는 RTLS에서 TWR 또는 TDoA 방식의 ranging/positioning에 활용될 수 있다[3]. "
        "본 연구에서는 anchor 좌표를 알고 있다고 가정하고, tag-anchor 거리 측정값으로 TurtleBot의 2D 위치를 추정한다."
    )
    for eq in EQUATIONS[:5]:
        add_equation(doc, eq["no"], eq["eq"], eq["desc"])
    add_figure(doc, figs.get("fig05_uwb_anchor_geometry.png", Path()), "Fig. 5. 4-anchor UWB 위치추정 geometry.", 6.2)

    doc.add_heading("5. 충전소 선택 및 LiDAR Local Approach", level=1)
    doc.add_paragraph(
        "추정된 UWB 위치는 두 충전소 중 가까운 station을 선택하는 데 사용된다. 선택된 station 좌표는 `/target_charger`로 publish되고, "
        "local planner는 `/scan`, `/uwb_pose`, `/target_charger`를 이용해 heading error를 줄이며 이동한다. "
        "front sector가 위험 거리 이하이면 회피 상태로 전환하고, 목표 반경 `R_goal=1.0 m`에 들어오면 정지한다."
    )
    for eq in EQUATIONS[5:9]:
        add_equation(doc, eq["no"], eq["eq"], eq["desc"])

    doc.add_heading("6. QR 기반 최종 정렬 모델", level=1)
    doc.add_paragraph(
        "OpenCV QRCodeDetector는 영상에서 QR code를 검출하고 decode하는 API를 제공한다[5]. 본 연구에서는 UWB-LiDAR 접근이 완료된 후 "
        "QR marker의 네 꼭짓점으로부터 marker 중심을 계산하고, 영상 중심과의 오차를 최종 docking 제어에 사용할 수 있다."
    )
    add_equation(doc, EQUATIONS[9]["no"], EQUATIONS[9]["eq"], EQUATIONS[9]["desc"])

    doc.add_heading("7. 예비 로그 분석 결과", level=1)
    doc.add_paragraph(
        "현재 확보된 실제 로그는 UWB 좌표 로그와 COM5 serial diagnostic log이다. true pose가 실시간으로 동기화된 최종 rosbag은 아직 없으므로, "
        "아래 결과는 논문 최종 성능값이 아니라 실험 설계를 위한 예비 분석값으로 제시한다."
    )
    add_table(
        doc,
        ["지표", "값", "해석"],
        [
            ["sample count", f"{metrics['n']:.0f}", "안정 정적 구간에서 선택한 UWB 표본 수"],
            ["mean error", f"{metrics['mean_error_cm']:.2f} cm", "중앙 기준점 (2.50, 2.50 m) 대비 평균 오차"],
            ["RMSE", f"{metrics['rmse_cm']:.2f} cm", "큰 오차에 민감한 종합 지표"],
            ["95% error", f"{metrics['p95_cm']:.2f} cm", "표본 95%가 포함되는 오차 수준"],
            ["x/y std", f"{metrics['x_std_cm']:.2f} / {metrics['y_std_cm']:.2f} cm", "필터링 후 반복성"],
        ],
        [1.4, 1.5, 4.1],
    )
    add_figure(doc, figs.get("fig07_uwb_static_scatter.png", Path()), "Fig. 6. 실제 UWB 로그 기반 정적 좌표 산점도.", 5.7)
    add_figure(doc, figs.get("fig08_uwb_static_error.png", Path()), "Fig. 7. 중앙 기준점 대비 UWB 정적 위치 오차.", 5.7)
    add_figure(doc, figs.get("fig09_all_uwb_log_trace.png", Path()), "Fig. 8. 저장된 UWB 좌표 로그 궤적.", 5.7)

    doc.add_heading("8. 최종 논문 실험 계획", level=1)
    add_table(
        doc,
        ["실험", "측정값", "논문 결과 표기"],
        [
            ["UWB 정적 정확도", "true_x, true_y, uwb_x, uwb_y", "mean error, RMSE, 95% error"],
            ["충전소 선택", "start pose, selected charger, expected charger", "selection success rate"],
            ["LiDAR 접근", "rosbag, state transition, cmd_vel, near_charger time", "approach time, stop radius error"],
            ["QR 최종 정렬", "QR pixel error, final physical offset", "final docking error cm"],
            ["통합 충전", "battery trigger to charging start", "success rate, elapsed time"],
        ],
        [1.45, 2.65, 2.9],
    )

    doc.add_heading("9. 결론", level=1)
    doc.add_paragraph(
        "본 초안은 UWB-LiDAR 기반 TurtleBot 충전소 접근 시스템을 논문 형식으로 정리한 것이다. "
        "현재 자료만으로는 UWB 예비 로그 기반 위치 안정성, 충전소 선택 구조, LiDAR local planner 구조까지는 논문 방법론으로 정리 가능하다. "
        "다만 최종 QR docking 및 실제 무선충전 성공률은 아직 실험 증거가 없으므로, 최종 논문에서는 반드시 true pose가 포함된 반복 실험과 "
        "최종 docking 사진/영상/rosbag을 추가해야 한다."
    )

    doc.add_heading("참고문헌", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.add_run(f"{ref['id']} {ref['title']}. ").bold = True
        p.add_run(ref["url"])

    doc.save(DOCX_PATH)


def write_csvs() -> None:
    with EQUATION_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["no", "name", "equation", "description"])
        for eq in EQUATIONS:
            writer.writerow([eq["no"], eq["name"], eq["eq"], eq["desc"]])
    with REF_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["id", "title", "url", "use"])
        for ref in REFERENCES:
            writer.writerow([ref["id"], ref["title"], ref["url"], ref["use"]])


def write_message(metrics: dict[str, float]) -> None:
    MSG_PATH.write_text(
        f"""교수님, 안녕하세요.

UWB TurtleBot 자율 무선충전 프로젝트를 실제 논문 초안 형태로 다시 정리했습니다.

이번 파일에는 단순 참고자료 정리가 아니라 초록, 서론, 관련 연구, 시스템 구성, UWB 위치추정 수식, 충전소 선택식, LiDAR local planner 제어식, QR 최종 정렬식, 실제 UWB 로그 기반 예비 결과, 최종 실험 계획, 참고문헌을 포함했습니다.

현재 실제 로그 기반 예비 분석에서는 중앙 기준점 기준 평균 오차 {metrics['mean_error_cm']:.2f} cm, RMSE {metrics['rmse_cm']:.2f} cm가 나왔습니다. 다만 true pose가 동기화된 최종 실험값은 아니므로, 본문에도 예비 분석값으로 명확히 표시했습니다.

논문 최종본으로 가려면 다음 실험에서 true_x/true_y가 포함된 UWB 반복 측정, rosbag 기반 접근 시간, QR 최종 정렬 오차, 실제 충전 시작 성공률을 추가하면 됩니다.

감사합니다.
""",
        encoding="utf-8",
    )


def main() -> None:
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figs = copy_figures()
    metrics = compute_metrics()
    build_docx(figs, metrics)
    write_csvs()
    write_message(metrics)
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    shutil.make_archive(str(ZIP_PATH.with_suffix("")), "zip", OUT_DIR)
    print(f"Wrote {OUT_DIR}")
    print(f"Wrote {ZIP_PATH}")


if __name__ == "__main__":
    main()
