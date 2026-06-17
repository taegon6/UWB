from __future__ import annotations

import csv
import math
import shutil
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "UWB_TurtleBot_교수님전달용_LiDAR파트_논문형_흑백_간소본"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "01_LiDAR파트_연구정리_논문형_흑백_간소본.docx"
CSV_PATH = OUT_DIR / "02_LiDAR파트_수식_참고문헌_색인.csv"
ZIP_PATH = ROOT / "UWB_TurtleBot_교수님전달용_LiDAR파트_논문형_흑백_간소본.zip"


BLACK = "#111111"
DARK = "#333333"
MID = "#666666"
LIGHT = "#dddddd"
VERY_LIGHT = "#f5f5f5"


PARAMS = [
    ("emergency_dist", "0.30 m", "전방 대표거리가 이 값보다 작으면 선속도를 0으로 두고 즉시 회피 회전"),
    ("avoid_dist", "0.55 m", "장애물 회피 상태로 진입하는 전방 거리 조건"),
    ("clear_dist", "0.80 m", "회피 상태를 유지하다가 정상 추종으로 복귀하는 거리 조건"),
    ("min_valid_scan_range", "0.05 m", "비정상적으로 가까운 LaserScan 값을 제거하기 위한 하한"),
    ("max_linear", "0.08 m/s", "목표 방향 정렬이 충분할 때 사용하는 최대 선속도"),
    ("slow_linear", "0.04 m/s", "회피 중 또는 heading 오차가 중간 정도일 때 사용하는 저속 주행값"),
    ("max_angular", "0.60 rad/s", "목표 추종 및 회피 회전에 적용되는 각속도 상한"),
    ("kp_heading", "1.2", "목표 heading 오차에 대한 비례 제어 이득"),
    ("goal_radius", "1.00 m", "LiDAR 접근 구간을 종료하고 QR 정밀 도킹으로 넘기는 거리 기준"),
    ("scan_timeout", "0.5 s", "LiDAR 입력이 끊겼다고 판단하는 시간 기준"),
    ("uwb_timeout", "2.0 s", "UWB pose 입력이 끊겼다고 판단하는 시간 기준"),
]


EQUATIONS = [
    ("(1)", "좌표 변환", "p_c^base = T_base^map p_c^map", "충전소 목표점이 map/odom 기준으로 주어질 경우, LiDAR 기반 제어는 base frame에서 목표 방향을 계산한다."),
    ("(2)", "LaserScan 각도", "theta_i = angle_min + i * angle_increment", "ranges[i]가 로봇 전방 기준 어느 방향의 거리값인지 계산한다."),
    ("(3)", "유효 거리 집합", "R = {r_i | finite(r_i), max(range_min, r_min) < r_i <= range_max}", "NaN, Inf, 센서 하한 이하 및 상한 초과 값을 제어 계산에서 제외한다."),
    ("(4)", "전방 섹터", "S_F = {r_i in R | -20 deg <= theta_i <= 20 deg}", "정면 충돌 위험을 판단하는 각도 범위이다."),
    ("(5)", "전방 대표거리", "d_F = (1/K) sum_{k=1}^{K} r_(k),  K=max(1, floor(0.1N))", "전방 거리값을 오름차순 정렬한 뒤 하위 10% 평균을 사용한다."),
    ("(6)", "좌우 절사평균", "d_side = mean({r_(k) | 0.2N <= k <= 0.8N})", "측면 여유 공간은 20-80% 절사평균으로 계산한다."),
    ("(7)", "회피 점수", "J_L = 0.7d_FL + 0.3d_L,  J_R = 0.7d_FR + 0.3d_R", "전방 좌우 공간을 더 크게 반영하여 회피 방향을 선택한다."),
    ("(8)", "목표 거리", "d_goal = sqrt((x_c-x_r)^2 + (y_c-y_r)^2)", "UWB pose와 선택된 충전소 목표점 사이의 평면 거리이다."),
    ("(9)", "목표 방향 오차", "e_theta = wrap(atan2(y_c-y_r, x_c-x_r) - theta_r)", "로봇 heading과 목표점 방향 사이의 각도 오차이다."),
    ("(10)", "목표 추종 각속도", "omega_goal = clamp(k_p e_theta, -omega_max, omega_max)", "장애물이 없을 때 목표 방향으로 회전하기 위한 비례 제어식이다."),
    ("(11)", "목표 추종 선속도", "v = 0 if |e_theta|>20 deg; v_slow if 10 deg<|e_theta|<=20 deg; v_max otherwise", "목표 방향과 크게 틀어진 상태에서 전진하지 않도록 한다."),
    ("(12)", "회피 진입", "EMERGENCY if d_F<0.30; AVOID if d_F<0.55; CLEAR if d_F>=0.80", "코드의 세 거리 임계값을 상태 전이 조건으로 정리한 식이다."),
    ("(13)", "회피 명령", "v_cmd=v_slow, omega_cmd=+0.7omega_max if J_L>J_R else -0.7omega_max", "좌우 여유 공간 점수에 따라 회피 회전 방향을 선택한다."),
    ("(14)", "비상 정지 회전", "v_cmd=0, omega_cmd=+omega_max if J_L>J_R else -omega_max", "전방 거리가 emergency_dist보다 작으면 전진을 멈춘 뒤 큰 여유 공간 쪽으로 회전한다."),
    ("(15)", "근접 handoff", "near_charger = true if d_goal < 1.00 m", "LiDAR 근거리 접근이 끝나면 QR 기반 최종 정렬 단계로 넘긴다."),
    ("(16)", "권장 보완식", "v_safe <= sqrt(2 a_brake max(d_F-d_safe,0))", "현재 구현을 논문화할 때 추가할 수 있는 제동거리 기반 안전속도 상한이다. 실제 코드에는 아직 별도 파라미터로 구현되어 있지 않다."),
]


REFERENCES = [
    ("[1]", "ROS sensor_msgs/LaserScan message documentation", "https://docs.ros.org/en/kinetic/api/sensor_msgs/html/msg/LaserScan.html"),
    ("[2]", "ROS tf package documentation", "https://docs.ros.org/en/kinetic/api/tf/html/c++/"),
    ("[3]", "J. Borenstein and Y. Koren, The Vector Field Histogram - Fast Obstacle Avoidance for Mobile Robots, IEEE Transactions on Robotics and Automation, 1991.", "https://public.websites.umich.edu/~ykoren/uploads/The_Vector_Field_HistogramuFast_Obstacle_Avoidance.pdf"),
    ("[4]", "D. Fox, W. Burgard, and S. Thrun, The Dynamic Window Approach to Collision Avoidance, IEEE Robotics & Automation Magazine, 1997.", "https://cir.nii.ac.jp/crid/1363388843526965632"),
]


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\malgunbd.ttf" if bold else r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "777777"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill="E8E8E8"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "Malgun Gothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(60, 60, 60)


def add_figure(doc, path, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def draw_arrow(draw, start, end, width=3):
    draw.line([start, end], fill=BLACK, width=width)
    sx, sy = start
    ex, ey = end
    angle = math.atan2(ey - sy, ex - sx)
    size = 13
    pts = [
        (ex, ey),
        (ex - size * math.cos(angle - 0.45), ey - size * math.sin(angle - 0.45)),
        (ex - size * math.cos(angle + 0.45), ey - size * math.sin(angle + 0.45)),
    ]
    draw.polygon(pts, fill=BLACK)


def draw_box(draw, box, title, body, fill=VERY_LIGHT):
    draw.rectangle(box, fill=fill, outline=BLACK, width=2)
    x1, y1, _, _ = box
    draw.text((x1 + 22, y1 + 18), title, font=font(26, True), fill=BLACK)
    for idx, line in enumerate(body):
        draw.text((x1 + 22, y1 + 58 + idx * 34), line, font=font(21), fill=DARK)


def make_io_figure(path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 45), "ROS topic interface of the LiDAR local planner", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "The module boundary is defined by input topics, command output, state output, and QR handoff signal.", font=font(21), fill=DARK)
    planner = (590, 315, 980, 565)
    draw.rectangle(planner, fill="white", outline=BLACK, width=4)
    draw.text((650, 390), "LiDAR Local Planner", font=font(29, True), fill=BLACK)
    draw.text((660, 440), "sector statistics\nstate machine\nvelocity command", font=font(20), fill=DARK, spacing=8)
    inputs = [
        ("/scan\nLaserScan", (110, 190, 430, 290), (590, 355)),
        ("/uwb_pose\nPose2D", (110, 330, 430, 430), (590, 410)),
        ("/target_charger\nPose2D", (110, 470, 430, 570), (590, 465)),
        ("/uwb_pose_status\nString", (110, 610, 430, 710), (590, 520)),
    ]
    for label, box, end in inputs:
        draw.rectangle(box, fill=VERY_LIGHT, outline=BLACK, width=2)
        draw.multiline_text((box[0] + 22, box[1] + 22), label, font=font(22, True), fill=BLACK, spacing=6)
        draw_arrow(draw, (box[2], (box[1] + box[3]) // 2), end)
    outputs = [
        ("/cmd_vel\nTwist", (1160, 260, 1410, 350), (980, 390)),
        ("/lidar_state\nString", (1160, 430, 1410, 520), (980, 465)),
        ("/near_charger\nBool", (1160, 600, 1410, 690), (980, 530)),
    ]
    for label, box, start in outputs:
        draw.rectangle(box, fill=VERY_LIGHT, outline=BLACK, width=2)
        draw.multiline_text((box[0] + 22, box[1] + 20), label, font=font(22, True), fill=BLACK, spacing=6)
        draw_arrow(draw, start, (box[0], (box[1] + box[3]) // 2))
    img.save(path)


def make_frame_figure(path):
    img = Image.new("RGB", (1500, 820), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 45), "Frame conversion used before LiDAR local planning", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "Target point and LaserScan sectors must be interpreted in a consistent robot-centered frame.", font=font(21), fill=DARK)
    draw_box(draw, (90, 250, 430, 410), "map / odom", ["UWB pose", "charger target"])
    draw_box(draw, (595, 250, 930, 410), "TF transform", ["p_c^base = T_base^map p_c^map"], fill="white")
    draw_box(draw, (1095, 250, 1435, 410), "base frame", ["heading error", "LaserScan sectors"])
    draw_arrow(draw, (430, 330), (595, 330))
    draw_arrow(draw, (930, 330), (1095, 330))
    draw.rectangle((210, 570, 1290, 705), fill="white", outline=BLACK, width=2)
    draw.text((245, 600), "Interpretation", font=font(25, True), fill=BLACK)
    draw.text((470, 592), "LiDAR detects nearby obstacles in the robot base frame.", font=font(21), fill=DARK)
    draw.text((470, 632), "The UWB target is converted to the same frame before heading control.", font=font(21), fill=DARK)
    img.save(path)


def make_sector_figure(path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    cx, cy = 750, 560
    radius = 300
    draw.text((60, 45), "LaserScan sector definition and robust distance metric", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "The implemented planner uses five angular sectors and conservative front-distance estimation.", font=font(21), fill=DARK)
    for deg in (-120, -70, -20, 20, 70, 120):
        rad = math.radians(deg)
        end = (cx + radius * math.cos(rad), cy - radius * math.sin(rad))
        draw.line((cx, cy, end[0], end[1]), fill=LIGHT, width=2)
        draw.text((end[0] - 22, end[1] - 18), f"{deg}deg", font=font(16), fill=MID)
    draw.arc((cx - radius, cy - radius, cx + radius, cy + radius), start=200, end=340, fill=BLACK, width=3)
    draw.polygon([(cx, cy - 23), (cx - 16, cy + 20), (cx + 16, cy + 20)], outline=BLACK, fill="white")
    draw.text((cx - 28, cy + 30), "robot", font=font(17), fill=BLACK)
    labels = [
        ("left", 104, 375),
        ("front-left", 45, 270),
        ("front\nlowest 10%", 0, 210),
        ("front-right", -45, 270),
        ("right", -104, 375),
    ]
    for label, deg, rr in labels:
        rad = math.radians(deg)
        x = cx + rr * math.cos(rad)
        y = cy - rr * math.sin(rad)
        draw.text((x - 70, y - 18), label, font=font(20, True), fill=BLACK)
    draw.rectangle((70, 665, 575, 815), outline=BLACK, width=2, fill="white")
    draw.text((95, 692), "Front distance: mean of the lowest 10% samples", font=font(20), fill=BLACK)
    draw.text((95, 730), "Side distance: 20-80% trimmed mean", font=font(20), fill=BLACK)
    draw.text((95, 768), "Invalid values: NaN, Inf, out-of-range samples", font=font(20), fill=BLACK)
    img.save(path)


def make_metric_flow_figure(path):
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 45), "Distance representative calculation from LaserScan sectors", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "The variable front_min is implemented as a lower-10% mean, not as a single minimum sample.", font=font(21), fill=DARK)
    boxes = [
        ((90, 250, 340, 350), "LaserScan\nranges"),
        ((470, 250, 760, 350), "valid range\nfiltering"),
        ((890, 170, 1240, 300), "front sector\nsort -> lower 10% mean\n= d_F"),
        ((890, 430, 1240, 560), "side sectors\nsort -> 20-80% mean\n= d_FL,d_FR,d_L,d_R"),
    ]
    for box, label in boxes:
        draw.rectangle(box, fill="white", outline=BLACK, width=2)
        draw.multiline_text((box[0] + 22, box[1] + 22), label, font=font(22, True), fill=BLACK, spacing=6)
    draw_arrow(draw, (340, 300), (470, 300))
    draw_arrow(draw, (760, 300), (890, 235))
    draw_arrow(draw, (760, 325), (890, 495))
    draw.rectangle((165, 685, 1335, 795), fill=VERY_LIGHT, outline=BLACK, width=2)
    draw.text((205, 718), "Reason for notation:", font=font(23, True), fill=BLACK)
    draw.text((470, 716), "front_min follows the code variable name, but the value is a conservative statistic.", font=font(21), fill=DARK)
    draw.text((470, 756), "This distinction should be written explicitly in the report.", font=font(21), fill=DARK)
    img.save(path)


def make_state_figure(path):
    img = Image.new("RGB", (1500, 930), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 45), "LiDAR local planner state transition", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "The controller is implemented as a small reactive state machine rather than a global planner.", font=font(21), fill=DARK)
    boxes = {
        "WAIT": (95, 220, 345, 330),
        "STALE_INPUT": (95, 450, 345, 560),
        "GO_TO_TARGET": (560, 220, 840, 330),
        "AVOID": (560, 450, 840, 560),
        "EMERGENCY_STOP": (1070, 450, 1370, 560),
        "NEAR_CHARGER": (1070, 220, 1370, 330),
    }
    for label, box in boxes.items():
        draw.rectangle(box, fill="white", outline=BLACK, width=2)
        x1, y1, x2, y2 = box
        text = "AVOID_LEFT /\nAVOID_RIGHT" if label == "AVOID" else label
        draw.multiline_text((x1 + 24, y1 + 32), text, font=font(24, True), fill=BLACK, spacing=6)
    draw_arrow(draw, (345, 275), (560, 275))
    draw_arrow(draw, (840, 275), (1070, 275))
    draw_arrow(draw, (700, 330), (700, 450))
    draw_arrow(draw, (840, 505), (1070, 505))
    draw_arrow(draw, (560, 505), (345, 505))
    draw.text((398, 240), "valid inputs", font=font(17), fill=DARK)
    draw.text((865, 240), "d_goal < 1.00 m", font=font(17), fill=DARK)
    draw.text((718, 372), "d_F < 0.55 m", font=font(17), fill=DARK)
    draw.text((870, 468), "d_F < 0.30 m", font=font(17), fill=DARK)
    draw.text((365, 470), "timeout or bad UWB status", font=font(17), fill=DARK)
    draw.rectangle((170, 690, 1330, 815), fill="white", outline=BLACK, width=2)
    draw.text((205, 720), "Hysteresis:", font=font(23, True), fill=BLACK)
    draw.text((370, 718), "avoidance starts at 0.55 m but returns to target tracking only after 0.80 m.", font=font(21), fill=DARK)
    draw.text((370, 758), "This reduces state chattering caused by noisy short-range scans.", font=font(21), fill=DARK)
    img.save(path)


def make_handoff_figure(path):
    img = Image.new("RGB", (1500, 850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 45), "QR docking handoff condition", font=font(34, True), fill=BLACK)
    draw.text((62, 92), "The LiDAR planner stops at the near-charger region and does not claim final docking accuracy.", font=font(21), fill=DARK)
    cx, cy = 1030, 440
    draw.ellipse((cx - 180, cy - 180, cx + 180, cy + 180), outline=BLACK, width=3)
    draw.ellipse((cx - 12, cy - 12, cx + 12, cy + 12), fill=BLACK)
    draw.text((cx - 62, cy - 230), "charger", font=font(24, True), fill=BLACK)
    draw.text((cx - 95, cy + 198), "goal_radius = 1.00 m", font=font(21), fill=DARK)
    path_points = [(165, 590), (310, 540), (455, 495), (610, 460), (770, 440), (850, 440)]
    draw.line(path_points, fill=BLACK, width=3)
    for pt in path_points[1:-1]:
        draw.ellipse((pt[0] - 5, pt[1] - 5, pt[0] + 5, pt[1] + 5), fill=BLACK)
    draw_arrow(draw, path_points[-2], path_points[-1])
    draw.rectangle((145, 635, 705, 735), fill=VERY_LIGHT, outline=BLACK, width=2)
    draw.text((175, 662), "LiDAR local planner region:", font=font(22, True), fill=BLACK)
    draw.text((175, 700), "low-speed approach and obstacle avoidance", font=font(20), fill=DARK)
    draw.rectangle((835, 635, 1345, 735), fill="white", outline=BLACK, width=2)
    draw.text((865, 662), "QR docking region:", font=font(22, True), fill=BLACK)
    draw.text((865, 700), "near_charger = true, then precision docking", font=font(20), fill=DARK)
    img.save(path)


def setup_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(20, 20, 20)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for name, size in (("Heading 1", 13), ("Heading 2", 11)):
        style = styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_run_font(run, size=None, bold=None):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TurtleBot 무선충전 주행을 위한 LiDAR 섹터 기반 Local Planner 설계 및 구현")
    add_run_font(r, 17, True)
    r.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UWB-TurtleBot 자율 무선충전 프로젝트 | 교수님 검토용 연구 정리")
    add_run_font(r, 9, False)
    r.font.color.rgb = RGBColor(70, 70, 70)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("작성 범위: LiDAR local planner 구현, 수식화, 가정 및 한계 정리")
    add_run_font(r, 8.5, False)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.15)
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.add_run("- " + text)
    return p


def build_doc(figs):
    doc = Document()
    setup_doc_styles(doc)
    add_title(doc)

    doc.add_heading("요약", level=1)
    add_paragraph(
        doc,
        "본 문서는 UWB 기반 TurtleBot 자율 무선충전 시스템에서 LiDAR local planner 모듈의 설계 및 구현 내용을 정리한다. "
        "해당 모듈은 /scan, /uwb_pose, /target_charger, /uwb_pose_status를 입력으로 사용하며, 장애물 회피 상태에 따라 /cmd_vel, /lidar_state, /near_charger를 출력한다. "
        "충전기 반경 1.00 m 이내에서는 LiDAR 주행 제어를 종료하고 QR 기반 정밀 도킹 단계로 제어권을 이양하도록 구성하였다."
    )

    doc.add_heading("1. 정리 범위", level=1)
    add_paragraph(
        doc,
        "본 문서는 UWB-TurtleBot 자율 무선충전 시스템 중 LiDAR가 담당하는 근거리 장애물 회피 및 충전소 접근 안정화 로직을 정리한 것이다. "
        "해당 모듈은 전역 경로계획이나 SLAM 지도를 생성하는 부분이 아니라, UWB가 제공하는 로봇 pose와 충전소 목표점을 바탕으로 저속 접근 중 전방 장애물에 반응하는 reactive local planner이다. "
        "따라서 본 정리는 실제 구현 파일 scripts/lidar_local_planner.py와 config/planner_params.yaml의 파라미터를 기준으로 작성하였다."
    )
    add_paragraph(
        doc,
        "LiDAR 파트의 역할은 세 가지로 구분된다. 첫째, /scan으로 들어오는 LaserScan 거리 배열을 각도 기준 섹터로 분리한다. 둘째, 전방 및 좌우 공간의 대표 거리를 계산하여 회피가 필요한 상태를 판단한다. 셋째, 목표 방향 추종 명령과 장애물 회피 명령 중 현재 상태에 맞는 /cmd_vel을 출력한다."
    )

    doc.add_heading("2. 시스템 입출력과 좌표계 가정", level=1)
    add_figure(doc, figs["io"], "Fig. 1. UWB-TurtleBot 시스템에서 LiDAR local planner의 ROS topic 입출력 구조.", 6.15)
    add_table(
        doc,
        ["구분", "ROS 인터페이스", "LiDAR planner에서의 의미"],
        [
            ["입력", "/scan (sensor_msgs/LaserScan)", "전방 및 좌우 장애물 거리 판단에 사용되는 2D planar scan"],
            ["입력", "/uwb_pose (geometry_msgs/Pose2D)", "UWB/odometry 기반 로봇 위치와 heading"],
            ["입력", "/target_charger (geometry_msgs/Pose2D)", "선택된 충전소 목표 좌표"],
            ["입력", "/uwb_pose_status (std_msgs/String)", "pose 품질이 허용 상태인지 확인하는 안전 gate"],
            ["출력", "/cmd_vel (geometry_msgs/Twist)", "선속도 및 각속도 명령"],
            ["출력", "/lidar_state, /near_charger", "현재 planner 상태와 QR 도킹 handoff 신호"],
        ],
        [0.9, 2.3, 3.8],
    )
    add_paragraph(
        doc,
        "LaserScan의 각도는 scan frame의 +x 방향을 0 rad로 두고 +z축 기준 반시계 방향으로 증가한다고 해석한다[1]. UWB pose 또는 충전소 좌표가 map/odom frame 기준이라면, 목표점은 제어 전에 base frame으로 변환되어야 한다. ROS tf는 시간에 따라 변하는 여러 좌표계 사이의 관계를 관리하고, 임의의 두 frame 사이에서 point 또는 vector를 변환하는 기능을 제공한다[2]."
    )
    add_figure(doc, figs["frame"], "Fig. 2. LiDAR local planner에서 사용하는 좌표계 변환 가정.", 6.1)

    doc.add_heading("3. LaserScan 전처리와 섹터 분할", level=1)
    add_paragraph(
        doc,
        "구현에서는 LaserScan의 각 range 값을 순회하면서 NaN, Inf, 센서 유효 범위를 벗어난 값을 제거한다. 이후 각 index의 방위각을 계산하고, 전방, 전방 좌측, 전방 우측, 좌측, 우측의 다섯 섹터로 분류한다. 전방 섹터는 -20 deg부터 +20 deg까지이며, 장애물과의 직접 충돌 가능성을 판단하는 가장 중요한 지표로 사용된다."
    )
    add_figure(doc, figs["sector"], "Fig. 3. LaserScan 섹터 분할 및 robust 거리 지표.", 6.1)

    doc.add_heading("4. Robust 거리 통계량", level=1)
    add_paragraph(
        doc,
        "전방 거리는 단순 최솟값을 그대로 사용하지 않고 정렬된 거리값 중 하위 10% 평균으로 계산하였다. 이 방식은 전방에 가까운 장애물이 있을 때 민감하게 반응하면서도, 단일 노이즈 값 하나가 전체 제어를 과도하게 지배하는 문제를 줄이기 위한 절충이다. 좌우 공간 평가는 20-80% 절사평균을 사용하여 양 끝단의 큰 오차나 순간적인 이상치를 줄인다."
    )
    add_paragraph(
        doc,
        "코드의 변수명은 front_min이지만, 실제 값은 단일 최소 거리값이 아니라 전방 섹터 유효 거리값의 하위 10% 평균이다. "
        "따라서 본 문서에서는 해당 값을 전방 대표거리 d_F로 표기하고, 변수명과 통계량의 차이를 구분하여 서술한다."
    )
    add_figure(doc, figs["metric"], "Fig. 4. LaserScan 섹터별 거리 대표값 계산 흐름.", 6.15)
    add_table(
        doc,
        ["거리량", "계산 방식", "사용 목적"],
        [
            ["d_F", "전방 섹터 하위 10% 평균", "비상정지와 회피 진입 판단"],
            ["d_FL, d_FR", "전방 좌우 섹터 절사평균", "회피 방향 점수의 주된 항"],
            ["d_L, d_R", "측면 섹터 절사평균", "회피 방향 점수의 보조 항"],
            ["J_L, J_R", "0.7 전방측 + 0.3 측면", "좌/우 회피 방향 선택"],
        ],
        [1.1, 2.7, 3.0],
    )

    doc.add_heading("5. 상태 기반 제어 구조", level=1)
    add_paragraph(
        doc,
        "제어기는 연속 최적화 기반 planner가 아니라, 입력 유효성, 목표점 거리, 전방 장애물 거리 조건에 따라 상태를 전환하는 reactive state machine으로 구현되어 있다. 입력이 없거나 오래된 경우에는 WAIT 또는 STALE_INPUT 상태에서 정지한다. 목표 반경 안으로 들어오면 NEAR_CHARGER를 publish하고, 이후 QR 기반 정밀 도킹 단계로 넘어갈 수 있도록 한다."
    )
    add_figure(doc, figs["state"], "Fig. 5. LiDAR local planner 상태 전이 구조.", 6.2)
    add_table(
        doc,
        ["상태", "진입 조건", "출력 명령"],
        [
            ["WAIT", "pose, target, scan 중 하나가 없음", "정지"],
            ["STALE_INPUT", "scan/pose timeout 또는 허용되지 않은 UWB status", "정지"],
            ["GO_TO_TARGET", "장애물이 없고 목표 반경 밖", "heading 비례 제어와 단계적 선속도"],
            ["AVOID_LEFT/RIGHT", "d_F < 0.55 m, 단 d_F >= 0.30 m", "저속 전진 + 여유 공간 방향 회전"],
            ["EMERGENCY_STOP", "d_F < 0.30 m", "선속도 0 + 여유 공간 방향 회전"],
            ["NEAR_CHARGER", "d_goal < 1.00 m 또는 latch 유지", "정지 + near_charger true"],
        ],
        [1.4, 3.1, 2.3],
    )

    doc.add_heading("6. 목표 추종 및 장애물 회피 제어식", level=1)
    add_paragraph(
        doc,
        "목표 추종 상태에서는 로봇 위치에서 충전소 목표점까지의 방향을 계산하고, 현재 heading과의 차이를 비례 제어 입력으로 사용한다. heading 오차가 20 deg보다 크면 회전만 수행하고, 10-20 deg 구간에서는 slow_linear로 제한하며, 10 deg 이하에서는 max_linear를 사용한다. 이 조건은 로봇이 목표 방향과 크게 어긋난 상태에서 전진하는 것을 방지한다."
    )
    add_paragraph(
        doc,
        "장애물 회피 상태에서는 왼쪽 회피 점수와 오른쪽 회피 점수를 비교한다. 왼쪽 공간이 더 넓으면 양의 각속도, 오른쪽 공간이 더 넓으면 음의 각속도를 사용한다. 비상정지 상태에서는 선속도를 0으로 두고 회전만 수행한다. 이는 VFH 계열의 거리 기반 회피 방향 선택 개념[3]을 단순화하여, 지도 없이 저속 주행 조건에서 동작하도록 구현한 형태로 볼 수 있다."
    )

    doc.add_heading("7. 수식 정리", level=1)
    add_table(doc, ["번호", "항목", "수식", "문서 내 해석"], EQUATIONS, [0.55, 1.35, 2.55, 3.0])

    doc.add_heading("8. 구현 파라미터", level=1)
    add_paragraph(
        doc,
        "아래 값은 config/planner_params.yaml에 명시된 현재 실험 설정이다. 코드 내부의 기본값과 다를 수 있으므로, 실제 실행 설정을 기준으로 해석하였다."
    )
    add_table(doc, ["파라미터", "현재 설정값", "의미"], PARAMS, [1.7, 1.0, 4.8])

    doc.add_heading("9. QR 정밀 도킹과의 역할 분리", level=1)
    add_paragraph(
        doc,
        "LiDAR planner는 충전소 접점에 대한 최종 정렬을 직접 수행하지 않는다. 목표점까지의 거리가 goal_radius=1.00 m보다 작아지면 near_charger=true를 publish하고 정지한다. 이 신호는 QR 또는 카메라 기반의 최종 정밀 도킹 로직이 활성화될 수 있는 handoff 조건으로 해석된다. 따라서 LiDAR 파트의 성능 평가는 최종 충전 접점 정렬 정확도가 아니라, 목표 반경까지의 안전한 접근과 근거리 장애물 반응성에 초점을 두어야 한다."
    )
    add_figure(doc, figs["handoff"], "Fig. 6. goal_radius 기준 LiDAR 접근 구간과 QR 정밀 도킹 구간의 역할 분리.", 6.15)

    doc.add_heading("10. 한계 및 보완 방향", level=1)
    add_bullet(doc, "현재 구조는 reactive local planner이므로, 장애물 배치가 복잡한 환경에서는 local minimum 또는 반복 회피가 발생할 수 있다.")
    add_bullet(doc, "LaserScan은 2D 평면 거리 센서이므로, 센서 높이보다 높거나 낮은 장애물은 검출하지 못할 수 있다.")
    add_bullet(doc, "동적 장애물의 속도 예측은 포함되어 있지 않다. 사람 또는 이동체가 많은 환경에서는 별도 추적 항이 필요하다.")
    add_bullet(doc, "현재 속도 제한은 heading 오차와 상태에 의해 piecewise로 결정된다. 논문형 확장에서는 DWA의 제동거리 조건[4]처럼 d_F 기반 안전속도 상한을 추가하면 제어 근거가 더 명확해진다.")
    add_bullet(doc, "목표점과 LiDAR scan frame 사이의 TF 관계가 불명확하면 heading 오차 계산 자체가 잘못될 수 있다. 따라서 실험 로그에는 scan frame, base frame, odom/map frame의 관계를 함께 기록하는 것이 바람직하다.")

    doc.add_heading("11. 참고문헌", level=1)
    for key, title, url in REFERENCES:
        p = doc.add_paragraph()
        r = p.add_run(f"{key} {title}. {url}")
        add_run_font(r, 8.5, False)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("LiDAR part research note - UWB TurtleBot autonomous charging project")
    add_run_font(r, 8, False)
    r.font.color.rgb = RGBColor(90, 90, 90)

    doc.save(DOCX_PATH)


def write_csv():
    with CSV_PATH.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(["type", "id", "title", "content", "note_or_url"])
        for no, title, eq, desc in EQUATIONS:
            writer.writerow(["equation", no, title, eq, desc])
        for key, title, url in REFERENCES:
            writer.writerow(["reference", key, title, "", url])


def build():
    if OUT_DIR.exists():
        shutil.rmtree(OUT_DIR)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figs = {
        "io": FIG_DIR / "fig1_ros_topic_interface_bw.png",
        "frame": FIG_DIR / "fig1_frame_conversion_bw.png",
        "sector": FIG_DIR / "fig2_laserscan_sector_bw.png",
        "metric": FIG_DIR / "fig3_distance_metric_flow_bw.png",
        "state": FIG_DIR / "fig3_state_machine_bw.png",
        "handoff": FIG_DIR / "fig4_qr_handoff_bw.png",
    }
    make_io_figure(figs["io"])
    make_frame_figure(figs["frame"])
    make_sector_figure(figs["sector"])
    make_metric_flow_figure(figs["metric"])
    make_state_figure(figs["state"])
    make_handoff_figure(figs["handoff"])
    build_doc(figs)
    write_csv()
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file in OUT_DIR.rglob("*"):
            if file.is_file():
                zf.write(file, file.relative_to(OUT_DIR.parent))
    print(f"Wrote {OUT_DIR}")
    print(f"Wrote {ZIP_PATH}")


if __name__ == "__main__":
    build()
